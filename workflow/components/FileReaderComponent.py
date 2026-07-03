import json
from dataclasses import dataclass

import requests

from workflow.basic.Component import Component, ComponentParameter
from workflow.basic.Node import Batch, ValueTypeOfIODefinition, VariableType
from workflow.WorkflowContext import NodeIOData, WorkflowContext


@dataclass
class FileReaderComponentInputDefinition:
    value_type: ValueTypeOfIODefinition
    content: list[str] | str
    parameter_name: str = "INPUT"


@dataclass
class FileReaderComponentOutputDefinition:
    variable_name: str = "OUTPUT"
    variable_type: VariableType = VariableType.OBJECT.value
    description: str | None = None
    schema: 'FileReaderComponentOutputDefinition' = None


@dataclass
class FileReaderComponentParam(ComponentParameter):
    output_definition: FileReaderComponentOutputDefinition
    input_definition: FileReaderComponentInputDefinition
    is_batch: bool = False
    batch: list[Batch] | None = None


class FileReaderComponent(Component[FileReaderComponentParam]):
    def __init__(self, component_parameter: FileReaderComponentParam, node_id: str):
        self.name = "FileReaderComponent"
        self.node_id = node_id
        self.component_parameter: FileReaderComponentParam = component_parameter
        super().__init__(component_parameter, node_id, self.name)

    async def process(self, input_data: dict | None = None, context: WorkflowContext | None = None,
                      **kwargs) -> dict:
        file_path = ""
        if context is not None:
            if self.component_parameter.is_batch:
                output_list = []
                ref_node_id = self.component_parameter.input_definition.content[0]
                ref_name = self.component_parameter.input_definition.content[1]
                file_list = context.get(str(ref_node_id)).output_data.get(ref_name)
                for file in file_list:
                    file_name, file_content = process_file_object(file)
                    output_list.append({"fileName": file_name, "fileContent": file_content})
                context.set(str(self.node_id), NodeIOData(
                    output_data={self.component_parameter.output_definition.variable_name: output_list}))
            else:
                if self.component_parameter.input_definition.value_type == ValueTypeOfIODefinition.REF.value:
                    ref_node_id = self.component_parameter.input_definition.content[0]
                    ref_name = self.component_parameter.input_definition.content[1]
                    file = context.get(str(ref_node_id)).output_data.get(ref_name)

                    # 判断是否为UploadFile
                    if str(type(file)) == "<class 'starlette.datastructures.UploadFile'>":
                        print("file is UploadFile")
                        file_name, file_content = await read_upload_file(file)

                        context.set(str(self.node_id),
                                    NodeIOData(output_data={self.component_parameter.output_definition.variable_name: {
                                        "fileName": file_name,
                                        "fileContent": file_content}}))

    def validate_inputs(self):
        pass

    def get_output_schema(self):
        pass

    @staticmethod
    def decode(json: json) -> 'FileReaderComponent':
        node_json = json['node']
        node_id = node_json['id']
        component_param = node_json['data']['componentParam']
        is_batch = component_param.get('isBatch', False)
        # batch = Batch.parse_batch(component_param.get('batch', None))
        input_definition = component_param['input_definition']
        parameter_name = input_definition[0]['parameter_name']
        value_type = input_definition[0]['value_type']
        content = input_definition[0]['content']
        FileReaderComponentInputDefinition(value_type=value_type, content=content, parameter_name=parameter_name)

        output_definition = component_param['output_definition']
        variable_name = output_definition['variable_name']
        variable_type = output_definition['variable_type']
        description = output_definition['description']
        schema = output_definition['schema']

        FileReaderComponentOutputDefinition(variable_name=variable_name, variable_type=variable_type,
                                            description=description, schema=schema)

        return FileReaderComponent(FileReaderComponentParam(
            FileReaderComponentOutputDefinition(variable_name=variable_name, variable_type=variable_type,
                                                description=description, schema=schema),
            FileReaderComponentInputDefinition(value_type=value_type, content=content, parameter_name=parameter_name),
            is_batch=is_batch),
            node_id)


class FileReadError(Exception):
    """自定义异常类，用于文件读取错误"""
    pass


import io

import docx
from fastapi import UploadFile
from pypdf import PdfReader


async def read_upload_file(file: UploadFile):
    contents = await file.read()
    file_extension = file.filename.split('.')[-1].lower()

    if file_extension == 'txt':
        try:
            decoded_contents = contents.decode('utf-8')
        except UnicodeDecodeError:
            decoded_contents = contents.decode('utf-8', errors='ignore')
    elif file_extension == 'docx':
        doc = docx.Document(io.BytesIO(contents))
        decoded_contents = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension == 'doc':
        try:
            content = contents  # 使用之前读取的内容，而不是再次读取
            print(f"File size: {len(content)} bytes")

            if len(content) == 0:
                return {"message": "The uploaded file is empty"}

            files = {"file": (file.filename, io.BytesIO(content), file.content_type)}

            response = requests.post("http://localhost:8080/api/doc/read", files=files)

            if response.status_code == 200:
                decoded_contents = response.text.replace('\n', '')
            else:
                return {"message": f"Error from Spring Boot: {response.status_code}", "details": response.text}

        except Exception as e:
            return {"message": f"An error occurred: {e!s}"}
    elif file_extension == 'pdf':
        pdf = PdfReader(io.BytesIO(contents))
        decoded_contents = '\n'.join([page.extract_text() for page in pdf.pages])
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

    return file.filename, decoded_contents


def process_file_object(file_object):
    """
    处理文件对象,提取文件名和内容。支持txt, docx, doc, pdf格式。

    参数:
    file_object (dict): 包含文件信息的字典

    返回:
    tuple: 包含文件名和解码后内容的元组
    """
    file_name = file_object.get('file_name')
    content = file_object.get('file_content').getvalue()
    file_extension = file_name.split('.')[-1].lower()

    if content is None:
        return file_name, ""

    if file_extension == 'txt':
        try:
            decoded_content = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                decoded_content = content.decode('gbk')
            except UnicodeDecodeError:
                decoded_content = content.decode('utf-8', errors='ignore')
    elif file_extension == 'docx':
        doc = docx.Document(io.BytesIO(content))
        decoded_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension == 'doc':
        try:
            print(f"File size: {len(content)} bytes")

            if len(content) == 0:
                return ""

            files = {"file": (file_name, io.BytesIO(content), 'application/msword')}

            response = requests.post("http://localhost:8080/api/doc/read", files=files)

            if response.status_code == 200:
                decoded_content = response.text.replace('\n', '')
            else:
                raise Exception(f"Error from Spring Boot: {response.status_code}, Details: {response.text}")

        except Exception as e:
            raise Exception(f"An error occurred: {e!s}")
    elif file_extension == 'pdf':
        pdf = PdfReader(io.BytesIO(content))
        decoded_content = '\n'.join([page.extract_text() for page in pdf.pages])
    else:
        raise Exception(f"Unsupported file type: {file_extension}")

    return file_name, decoded_content


async def read_file_object(file: UploadFile):
    async with file.file as f:
        # 使用 f 作为文件对象
        contents = await f.read()
    return contents


def get_file_info(file: UploadFile):
    return {
        "filename": file.filename,
        "content_type": file.content_type,
        "size": file.size
    }
