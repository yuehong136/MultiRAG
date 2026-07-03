import io
import os
from dataclasses import dataclass
from typing import Any

import chardet
from docx import Document
from fastapi import UploadFile

from workflow.utils.MinioOperator import MinioOperator
from workflow_v2.component.base_component import BaseComponent
from workflow_v2.utils import match_parameters
from workflow_v2.workflow_logging_config import WorkflowContextLogger


@dataclass
class BatchConfig:
    """批处理配置类"""

    batch_enable: bool = False
    batch_size: int = 100
    concurrent_size: int = 10
    input_lists: list[dict[str, Any]] = None

    @classmethod
    def from_batch_config(cls, config: dict[str, Any]) -> "BatchConfig":
        """从批处理配置创建实例"""
        if not config:
            return cls()

        input_lists = config.get("inputLists", [])
        if not isinstance(input_lists, list):
            input_lists = [input_lists]

        return cls(batch_enable=config.get("batchEnable", False), batch_size=config.get("batchSize", 100), concurrent_size=config.get("concurrentSize", 10), input_lists=input_lists)


class FileReaderComponent(BaseComponent):
    """LLM组件"""

    def __init__(self, component_id: str, title: str, node_data: dict[str, Any], logger: WorkflowContextLogger, **kwargs):
        super().__init__(component_id, title, logger)
        self.batch_config: BatchConfig = self._extract_batch_config(node_data)

        self.db = kwargs.get("db", None)
        self.user = kwargs.get("user", None)

    def _extract_batch_config(self, node_data: dict[str, Any]) -> BatchConfig:
        """从节点数据中提取批处理配置"""
        batch_data = node_data["data"]["inputs"].get("batch", {})
        return BatchConfig.from_batch_config(batch_data)

    async def execute(self) -> dict[str, Any]:
        self.logger.info(f"LLMComponent {self.title} execute")
        self.logger.info(f"LLMComponent {self.title} inputs: {self.inputs}")

        if self.batch_config.batch_enable:
            # 目前批量只支持minio文件
            output_list = []
            minio_file_info = next(iter(match_parameters(self.batch_config.input_lists, self.nodes).values()))
            bucket_name = minio_file_info.split("+")[0]
            dir_name = minio_file_info.split("+")[1]
            minio_operator = MinioOperator()
            file_list = minio_operator.list_objects(bucket_name=bucket_name, prefix=dir_name, recursive=False)
            for file in file_list:
                file_name = file["name"]
                file_data = minio_operator.download_to_memory(bucket_name=bucket_name, object_name=file_name)
                if file_name.endswith(".docx"):
                    file_content = read_docx_content(file_data)
                elif file_name.endswith(".txt"):
                    file_content = detect_and_read_content(file_data)
                else:
                    continue
                output_list.append({"fileName": file_name.split("/")[-1], "fileContent": file_content})

            return {"outputList": output_list}
        else:
            # 目前单文件只支持普通上传文件
            file_info = await parse_uploaded_file(next(iter(self.inputs.values())))
            file_name = file_info["file_name"]
            file_content = file_info["file_content"]
            return {"output": {"fileName": file_name, "fileContent": file_content}}

    async def execute_alone(self, input_value: dict, batch_value: dict | None = None) -> dict:
        self.logger.info(f"LLMComponent {self.title} execute")
        self.logger.info(f"LLMComponent {self.title} inputs: {self.inputs}")

        if self.batch_config.batch_enable:
            output_list = []
            minio_file_info = next(iter(batch_value.values()))
            bucket_name = minio_file_info.split("+")[0]
            dir_name = minio_file_info.split("+")[1]
            minio_operator = MinioOperator()
            file_list = minio_operator.list_objects(bucket_name=bucket_name, prefix=dir_name, recursive=False)
            for file in file_list:
                file_name = file["name"]
                file_data = minio_operator.download_to_memory(bucket_name=bucket_name, object_name=file_name)
                if file_name.endswith(".docx"):
                    file_content = read_docx_content(file_data)
                elif file_name.endswith(".txt"):
                    file_content = detect_and_read_content(file_data)
                else:
                    continue
                output_list.append({"fileName": file_name.split("/")[-1], "fileContent": file_content})

            return {"outputList": output_list}
        else:
            file_info = await parse_uploaded_file(next(iter(input_value.values())))
            file_name = file_info["file_name"]
            file_content = file_info["file_content"]
            return {"output": {"fileName": file_name, "fileContent": file_content}}


async def parse_uploaded_file(file: UploadFile) -> dict:
    """
    解析上传的文件，支持 docx 和 txt 格式

    Args:
        file (UploadFile): 上传的文件对象

    Returns:
        dict: 包含文件名和内容的字典
    """
    try:
        # 获取文件名和扩展名
        file_name = file.filename
        file_extension = os.path.splitext(file_name)[1].lower()

        # 读取文件内容
        content = await file.read()

        # 根据文件类型进行不同的处理
        if file_extension == ".docx":
            # 直接从内存读取，正文段落 + 表格单元格
            file_content = read_docx_content(io.BytesIO(content))

        elif file_extension == ".txt":
            # 直接解码文本内容
            file_content = content.decode("utf-8")

        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        return {"file_name": file_name, "file_content": file_content}

    except Exception as e:
        raise Exception(f"Error processing file: {e!s}")

    finally:
        # 关闭文件
        await file.close()


def read_docx_content(bytes_data) -> str:
    """
    从内存中的 docx(BytesIO)读取文本：正文段落 + 表格单元格。

    docx 是 zip(OOXML)包，不能按纯文本 decode；且表格内容不在 doc.paragraphs 里，
    必须单独遍历 doc.tables，否则内容大多在表格里的文档会读成空。
    """
    if bytes_data is None:
        return ""
    bytes_data.seek(0)
    doc = Document(bytes_data)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = "\t".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_bytes_content(bytes_data, encoding="utf-8", chunk_size=None):
    """
    从BytesIO对象中读取内容并转换为字符串

    参数:
        bytes_data: BytesIO对象
        encoding: 字符编码，默认utf-8
        chunk_size: 分块读取的大小，如果为None则一次性读取

    返回:
        str: 解码后的字符串内容
        None: 如果发生错误
    """
    try:
        if bytes_data is None:
            return None

        # 确保指针在开始位置
        bytes_data.seek(0)

        if chunk_size:
            # 分块读取
            content = ""
            while True:
                chunk = bytes_data.read(chunk_size)
                if not chunk:
                    break
                content += chunk.decode(encoding)
        else:
            # 一次性读取
            content = bytes_data.read().decode(encoding)

        return content

    except UnicodeDecodeError as e:
        print(f"解码错误，请检查编码格式是否正确: {e}")
        return None
    except Exception as e:
        print(f"读取内容时发生错误: {e}")
        return None
    finally:
        # 重置指针位置，以便后续可能的读取
        bytes_data.seek(0)


def detect_and_read_content(bytes_data, chunk_size=None):
    """
    自动检测编码并读取内容
    """
    try:
        if bytes_data is None:
            return None

        bytes_data.seek(0)
        raw_data = bytes_data.read()
        detected = chardet.detect(raw_data)
        encoding = detected["encoding"]

        # 重新封装成BytesIO
        new_bytes = io.BytesIO(raw_data)
        return read_bytes_content(new_bytes, encoding=encoding, chunk_size=chunk_size)

    except Exception as e:
        print(f"检测编码并读取内容时发生错误: {e}")
        return None
