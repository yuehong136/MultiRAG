import base64
import json
import logging
import os
import re
import shutil
import uuid
from datetime import datetime
from io import BytesIO

from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, UploadFile

from api.utils.api_utils import get_json_result

router = APIRouter()
logger = logging.getLogger(__name__)

# 临时文件存放目录
DOCX_LOG_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "temp", "docx_logs")
# 最大保留的请求数量
DOCX_LOG_MAX_REQUESTS = 10


def ensure_log_dir():
    """确保日志目录存在"""
    if not os.path.exists(DOCX_LOG_DIR):
        os.makedirs(DOCX_LOG_DIR)


def cleanup_old_logs():
    """
    清理旧的日志目录，只保留最近 DOCX_LOG_MAX_REQUESTS 次请求的数据
    按目录创建时间排序，删除最旧的
    """
    if not os.path.exists(DOCX_LOG_DIR):
        return

    # 获取所有子目录
    subdirs = []
    for name in os.listdir(DOCX_LOG_DIR):
        path = os.path.join(DOCX_LOG_DIR, name)
        if os.path.isdir(path):
            # 获取目录创建时间
            ctime = os.path.getctime(path)
            subdirs.append((path, ctime))

    # 如果目录数量未超过限制，不需要清理
    if len(subdirs) <= DOCX_LOG_MAX_REQUESTS:
        return

    # 按创建时间排序（最旧的在前）
    subdirs.sort(key=lambda x: x[1])

    # 删除最旧的目录，保留最近 DOCX_LOG_MAX_REQUESTS 个
    dirs_to_delete = subdirs[:-DOCX_LOG_MAX_REQUESTS]
    for dir_path, _ in dirs_to_delete:
        try:
            shutil.rmtree(dir_path)
            logger.info(f"[cleanup] 已删除旧日志目录: {os.path.basename(dir_path)}")
        except Exception as e:
            logger.warning(f"[cleanup] 删除目录失败 {dir_path}: {e}")


def save_docx_log(request_id: str, stage: str, file_content: bytes, json_data: dict = None, filename: str = None):
    """
    保存 docx 处理日志

    Args:
        request_id: 请求唯一标识符
        stage: 阶段标识，如 "input", "output"
        file_content: 文件内容（bytes）
        json_data: JSON 数据（可选）
        filename: 原始文件名（可选）
    """
    ensure_log_dir()

    # 在保存新日志前清理旧日志
    if stage == "input":
        cleanup_old_logs()

    # 创建请求专属目录
    request_dir = os.path.join(DOCX_LOG_DIR, request_id)
    if not os.path.exists(request_dir):
        os.makedirs(request_dir)

    # 保存文件
    file_suffix = f"_{filename}" if filename else ""
    docx_path = os.path.join(request_dir, f"{stage}{file_suffix}.docx")
    with open(docx_path, "wb") as f:
        f.write(file_content)

    # 保存 JSON 数据（如果有）
    if json_data is not None:
        json_path = os.path.join(request_dir, f"{stage}_data.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)

    return request_dir


def _replace_paragraph_text(paragraph, value: str):
    """
    更新单个段落的文本，同时保留段落级格式。
    仅替换文本内容；段落及其属性保持不变。
    """
    if value is None:
        value = ""

    runs = list(paragraph.runs)
    if runs:
        for extra_run in runs[1:]:
            extra_run._element.getparent().remove(extra_run._element)
        runs[0].text = value
    else:
        paragraph.add_run(value)


def update_cell_text_preserving_format(cell, text: str):
    """
    安全地更新单元格的文本内容，而不会破坏其段落结构或格式元数据（对齐、间距等）。
    """
    if text is None:
        text = ""

    lines = text.split('\n')
    paragraphs = list(cell.paragraphs)

    if not paragraphs:
        base_style = None
        new_paragraph = cell.add_paragraph('')
        paragraphs.append(new_paragraph)
    else:
        base_style = paragraphs[-1].style

    # Ensure we have enough paragraphs to accommodate all lines
    while len(paragraphs) < len(lines):
        new_paragraph = cell.add_paragraph('')
        if base_style:
            new_paragraph.style = base_style
        paragraphs.append(new_paragraph)

    # Update existing paragraphs with new text
    for idx, line in enumerate(lines):
        _replace_paragraph_text(paragraphs[idx], line)

    # Clear any remaining paragraphs to keep spacing while removing stale text
    for idx in range(len(lines), len(paragraphs)):
        _replace_paragraph_text(paragraphs[idx], "")


PLACEHOLDER_PATTERN = re.compile(r'\{\{([^{}]+)\}\}')

# 冒号模式（中英文冒号）
COLON_PATTERN = re.compile(r'[：:]')


def extract_underline_placeholders_from_paragraph(paragraph):
    """
    从段落中提取带下划线的填充项

    识别模式：标签文本 + 冒号 + 下划线区域（font.underline=True）
    例如："单   位：________________" 其中下划线部分的run的font.underline=True

    返回：列表，每个元素为 {
        'label': 标签文本,
        'underline_run_indices': 下划线run的索引列表,
        'colon_run_index': 冒号所在run的索引,
        'colon_position': 冒号在该run中的位置
    }
    """
    placeholders = []
    runs = list(paragraph.runs)

    if not runs:
        return placeholders

    # 构建run信息列表，记录每个run的文本、是否下划线、在段落中的字符起始位置
    run_info = []
    char_offset = 0
    for i, run in enumerate(runs):
        text = run.text or ""
        # 检查下划线属性
        is_underline = bool(run.font.underline)
        run_info.append({
            'index': i,
            'text': text,
            'is_underline': is_underline,
            'char_start': char_offset,
            'char_end': char_offset + len(text),
            'run': run
        })
        char_offset += len(text)

    # 找到所有下划线区域（连续的下划线run合并为一个区域）
    underline_regions = []
    i = 0
    while i < len(run_info):
        if run_info[i]['is_underline']:
            region_start = i
            region_end = i
            # 合并连续的下划线run
            while region_end + 1 < len(run_info) and run_info[region_end + 1]['is_underline']:
                region_end += 1
            underline_regions.append({
                'start_index': region_start,
                'end_index': region_end,
                'run_indices': list(range(region_start, region_end + 1))
            })
            i = region_end + 1
        else:
            i += 1

    # 对于每个下划线区域，向前查找标签
    for region in underline_regions:
        label = _find_label_before_underline(run_info, region['start_index'])
        if label:
            placeholders.append({
                'label': label['text'],
                'underline_run_indices': region['run_indices'],
                'colon_run_index': label['colon_run_index'],
                'colon_position': label['colon_position']
            })

    return placeholders


def _find_label_before_underline(run_info, underline_start_index):
    """
    从下划线区域向前查找标签

    规则：
    1. 向前查找最近的冒号（：或:）
    2. 冒号前的文本即为标签（去除前导空格，保留内部空格）
    3. 标签结束于上一个下划线区域或段落开头

    返回：{'text': 标签文本, 'colon_run_index': 冒号所在run索引, 'colon_position': 冒号位置}
    """
    # 收集下划线之前的所有文本和run信息
    text_before = ""
    run_text_mapping = []  # [(run_index, text_start_in_combined, text_end_in_combined)]

    for i in range(underline_start_index):
        text = run_info[i]['text']
        start_pos = len(text_before)
        text_before += text
        run_text_mapping.append((i, start_pos, start_pos + len(text)))

    if not text_before:
        return None

    # 查找最后一个冒号
    colon_match = None
    for match in COLON_PATTERN.finditer(text_before):
        colon_match = match

    if colon_match is None:
        return None

    colon_pos = colon_match.start()

    # 确定冒号所在的run
    colon_run_index = None
    colon_position_in_run = None
    for run_idx, start, end in run_text_mapping:
        if start <= colon_pos < end:
            colon_run_index = run_idx
            colon_position_in_run = colon_pos - start
            break

    if colon_run_index is None:
        return None

    # 提取冒号前的标签文本
    # 标签从上一个下划线区域结束后开始，或从段落开头开始
    label_start = 0
    for i in range(underline_start_index - 1, -1, -1):
        if run_info[i]['is_underline']:
            # 找到前一个下划线区域，标签从其后开始
            label_start = run_info[i]['char_end']
            break

    # 提取标签文本（冒号前的文本）
    label_text = text_before[label_start:colon_pos]

    # 清理标签文本
    # 去除前后空格，但保留内部空格（如"单   位"）
    label_text = label_text.strip()

    if not label_text:
        return None

    return {
        'text': label_text,
        'colon_run_index': colon_run_index,
        'colon_position': colon_position_in_run
    }


def process_paragraph_underline_placeholders(paragraph, placeholders_dict):
    """
    处理段落中的下划线填充项，生成占位符并更新段落

    Args:
        paragraph: docx段落对象
        placeholders_dict: 占位符字典，用于收集所有占位符

    Returns:
        bool: 是否有处理过占位符
    """
    underline_items = extract_underline_placeholders_from_paragraph(paragraph)

    if not underline_items:
        return False

    runs = list(paragraph.runs)
    processed = False

    for item in underline_items:
        label = item['label']
        underline_run_indices = item['underline_run_indices']

        # 生成标准化的占位符key
        # 移除多余空格，保留中文、数字、字母、下划线和常用符号（括号、顿号等）
        normalized_label = re.sub(r'\s+', '', label)  # 移除所有空格
        placeholder_key = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9_（）()、]', '', normalized_label)

        if not placeholder_key:
            continue

        placeholder_value = f"{{{{{placeholder_key}}}}}"

        # 计算原始下划线区域的总字符长度
        original_length = sum(len(runs[idx].text) for idx in underline_run_indices)

        # 如果占位符比原始长度短，居中填充空格以保持下划线长度
        if len(placeholder_value) < original_length:
            total_padding = original_length - len(placeholder_value)
            left_padding = total_padding // 2
            right_padding = total_padding - left_padding
            placeholder_value = ' ' * left_padding + placeholder_value + ' ' * right_padding

        # 将占位符文本填入第一个下划线run，清空其余下划线run
        if underline_run_indices:
            first_underline_run = runs[underline_run_indices[0]]
            first_underline_run.text = placeholder_value

            # 清空其余下划线run
            for idx in underline_run_indices[1:]:
                runs[idx].text = ""

        # 添加到占位符字典
        placeholders_dict[placeholder_key] = ""
        processed = True

    return processed


def fill_paragraph_underline_placeholders(paragraph, fill_data):
    """
    填充段落中的下划线占位符

    Args:
        paragraph: docx段落对象
        fill_data: 填充数据字典

    Returns:
        bool: 是否有填充过占位符
    """
    runs = list(paragraph.runs)
    filled = False

    for run in runs:
        if not run.text:
            continue

        original_length = len(run.text)

        # 检查run中是否有占位符
        def replace_placeholder(match):
            nonlocal filled
            key = match.group(1).strip()
            value = fill_data.get(key, "")
            if value is None:
                value = ""
            filled = True
            return str(value)

        new_text, count = PLACEHOLDER_PATTERN.subn(replace_placeholder, run.text)
        if count > 0:
            # 如果替换后的文本比原始文本短，居中填充空格以保持下划线长度
            if len(new_text) < original_length:
                total_padding = original_length - len(new_text)
                left_padding = total_padding // 2
                right_padding = total_padding - left_padding
                new_text = ' ' * left_padding + new_text + ' ' * right_padding
            run.text = new_text

    return filled


@router.post("/process_docx", summary="Word 文档占位符处理接口",
             response_description="返回处理后的文档和占位符 JSON 数据")
async def process_docx(
        file: UploadFile = File(...),
        data: str = Form(default=None)
):
    """
    ### POST `/v1/document/process_docx` Word 文档占位符处理接口

**功能描述**:
该接口接收用户上传的 Word 文件，对文件中的表格占位符进行解析、标准化和填充，最终生成处理后的文档与占位符的 JSON 数据。

---

### 请求体 (Request Body)

| 字段        | 类型          | 必填 | 描述                          |
|-------------|---------------|------|-------------------------------|
| `file`      | `UploadFile`  | 是   | 用户上传的 Word 文档，格式必须为 `.docx` |

---

### 响应 (Response)

#### 成功响应 (200)

- **响应格式**:
    - **`Content-Type: application/json`**
    - **示例**:
    ```json
    {
        "retcode": 0,
        "retmsg": "File processed successfully.",
        "data": {
            "placeholders": {
                "key1": "",
                "key2": ""
            },
            "file": "base64-encoded-processed-file-content"
        }
    }
    ```

- **字段说明**:
    | 字段             | 类型        | 描述                               |
    |------------------|-------------|------------------------------------|
    | `placeholders`   | `dict`      | 处理后的占位符键值对，键为占位符名称，值为空字符串 |
    | `file`           | `string`    | 处理后文档的 Base64 编码内容       |

---

#### 错误响应

| 状态码 | 错误类型              | 描述                                |
|--------|-----------------------|-------------------------------------|
| 400    | `Invalid file format` | 当上传的文件格式不是 `.docx` 时，返回此错误 |

- **示例**:
    ```json
    {
        "detail": "Invalid file format. Only .docx files are supported."
    }
    ```

---

### 主要流程

1. **文件校验**: 验证上传文件的格式是否为 `.docx`，如果不是，返回 400 错误。
2. **读取文档内容**: 使用 `python-docx` 读取上传文档中的表格内容。
3. **占位符处理**:
    - 判断和包装符合 `{{key}}` 格式的占位符。
    - 标准化占位符键名（仅保留中文字符、数字和下划线）。
    - 填充表格内容：根据规则右填充和下填充缺失内容。
    - 为下填充的占位符添加序号后缀。
4. **文档内容更新**: 将处理后的占位符填充回文档的对应表格。
5. **返回结果**:
    - 将处理后的文档内容以 Base64 格式返回。
    - 返回占位符的 JSON 数据，方便后续填写或二次处理。

---

### 注意事项

- **文件格式**: 仅支持 `.docx` 文件格式。
- **表格占位符规则**:
    - 占位符必须以 `{{` 开始，并以 `}}` 结束。
    - 键名仅支持中文字符和下划线，其他字符会被移除。
    - 自动为填充生成的占位符添加 `_1`、`_2` 等后缀，以避免重复。
- **返回文件编码**: 处理后的文档以 Base64 编码返回，需客户端自行解码并保存为 `.docx` 文件。

---
    """
    # 生成请求唯一标识符
    request_id = f"process_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
    logger.info(f"[process_docx] 请求ID: {request_id}, 文件名: {file.filename}")

    # 验证文件名和格式
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Invalid file format. Only .docx files are supported.")

    # 解析可选的 JSON 数据
    mapping_data = None
    if data:
        try:
            mapping_data = json.loads(data)
            logger.info(f"[process_docx] 输入 mapping_data: {json.dumps(mapping_data, ensure_ascii=False)}")
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid JSON data provided.")

    # 读取上传的文件内容
    input_file_content = await file.read()

    # 保存输入文件和 JSON 数据
    save_docx_log(request_id, "input", input_file_content, mapping_data, file.filename)
    logger.info("[process_docx] 输入文件已保存")

    # 工具方法：判断是否为独立单元格
    def is_isolated_cell(matrix, row_idx, col_idx):
        rows = len(matrix)
        cols = len(matrix[0]) if rows > 0 else 0

        # 检查上、下、左、右单元格是否有内容
        if row_idx > 0 and matrix[row_idx - 1][col_idx].strip():  # 上
            return False
        if row_idx < rows - 1 and matrix[row_idx + 1][col_idx].strip():  # 下
            return False
        if col_idx > 0 and matrix[row_idx][col_idx - 1].strip():  # 左
            return False
        if col_idx < cols - 1 and matrix[row_idx][col_idx + 1].strip():  # 右
            return False

        # 如果四个方向都没有内容，则为独立单元格
        return True

    # 工具方法：获取表格上方段落的描述
    def get_table_above_paragraphs(input_doc, table_idx):
        # 边界检查
        if table_idx < 0 or table_idx >= len(input_doc.tables):
            return []

        target_table = input_doc.tables[table_idx]
        tbl_elm = target_table._element
        parent = tbl_elm.getparent()
        if parent is None:
            return []

        # 在父容器中找到该表格的块级位置
        children = list(parent.iterchildren())
        try:
            pos = children.index(tbl_elm)
        except ValueError:
            return []

        # 建立一个从段落底层元素到其文本的映射，加速查找
        para_map = {id(p._element): p.text.strip() for p in input_doc.paragraphs}

        above_paragraphs = []
        # 向上回溯收集非空段落文本
        for j in range(pos - 1, -1, -1):
            child = children[j]
            # w:p 段落
            if child.tag.endswith('}p'):
                text = para_map.get(id(child), '')
                if text:
                    above_paragraphs.append(text)

        # 还原为阅读顺序（上->下）
        return above_paragraphs[::-1]

    # 工具方法：判断是否是占位符格式
    def is_placeholder_format(value: str) -> bool:
        return value.strip().startswith('{{') and value.strip().endswith('}}')

    # 工具方法：包装占位符
    def wrap_placeholder(value: str) -> str:
        val = value.strip()
        return val if is_placeholder_format(val) else f"{{{{{val}}}}}"

    # 工具方法：标准化占位符键名
    def normalize_placeholder_key(key: str) -> str:
        # 保留中文字符、数字和下划线
        return re.sub(r'[^\u4e00-\u9fa5_\d]', '', key)

    # 工具方法：生成JSON合法的完整键名
    def generate_full_placeholder_key(top_content: str) -> str:
        """
        基于完整的顶部内容生成JSON合法的占位符键名
        处理换行符、引号等特殊字符，确保JSON合法性
        """
        if not top_content or not top_content.strip():
            return ""

        # 1. 移除或替换冒号（通常在内容末尾）
        content = top_content.replace("：", "").replace(":", "")

        # 2. 将双引号替换为单引号
        content = content.replace('"', "'")

        # 3. 将换行符替换为分隔符，保持可读性
        content = content.replace('\n', ' | ')

        # 4. 清理多余的空格
        content = re.sub(r'\s+', ' ', content).strip()

        # 5. 移除其他可能导致JSON问题的字符
        # 保留中文、英文、数字、常用标点符号
        content = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s\'\(\)\[\]\{\}，。！？；、|（）【】《》]', '', content)

        # 6. 控制长度，避免键名过长（保留前200个字符）
        if len(content) > 200:
            content = content[:200] + "..."

        # 7. 确保不为空
        if not content.strip():
            return "未命名内容"

        return content.strip()

    # 工具方法：判断非空单元格是否需要处理
    def should_process_non_empty_cell(cell_text: str) -> bool:
        """
        判断非空单元格是否需要添加占位符
        条件：包含"："且不包含排除关键词
        """
        if not cell_text or not cell_text.strip():
            return False

        # 必须包含冒号
        if "：" not in cell_text:
            return False

        # 排除关键词列表
        exclude_keywords = ["√", "负责人签字", "盖章"]

        # 如果包含任何排除关键词，则不处理
        for keyword in exclude_keywords:
            if keyword in cell_text:
                return False

        return True

    # 工具方法：提取单元格顶部内容
    def extract_top_content(cell_text: str) -> str:
        """
        提取单元格的顶部内容
        通过换行符判断，取第一部分作为顶部内容
        """
        if not cell_text or not cell_text.strip():
            return ""

        # 按换行符分割
        lines = cell_text.split('\n')

        # 找到连续空行的位置，连续空行前的内容作为顶部
        top_content_lines = []
        consecutive_empty_count = 0

        for line in lines:
            if line.strip() == "":
                consecutive_empty_count += 1
                # 如果连续遇到2个或以上空行，认为顶部内容结束
                if consecutive_empty_count >= 2:
                    break
            else:
                # 遇到非空行，重置计数并添加到顶部内容
                consecutive_empty_count = 0
                top_content_lines.append(line)

        # 如果没有找到连续空行，说明整个内容都是顶部内容
        if not top_content_lines:
            top_content_lines = [line for line in lines if line.strip()]

        return '\n'.join(top_content_lines).strip()

    # 工具方法：标准化所有占位符
    def normalize_all_placeholders(matrix):
        key_map = {}
        rows = len(matrix)
        if rows == 0:
            return matrix, key_map

        cols = len(matrix[0])
        for r in range(rows):
            for c in range(cols):
                val = matrix[r][c].strip()
                if is_placeholder_format(val):
                    raw_key = val.strip('{').strip('}')
                    new_key = normalize_placeholder_key(raw_key)
                    key_map[raw_key] = new_key
                    matrix[r][c] = f"{{{{{new_key}}}}}"
        return matrix, key_map

    # 工具方法：判断单元格是否为合并单元格
    def is_merged_cell(cell):
        tc = cell._element
        vmerge = tc.xpath(".//w:vMerge")
        return vmerge and vmerge[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") != "restart"

    # 工具方法：判断单元格是否为下合并单元格的开始
    def is_start_of_down_merge(cell):
        tc = cell._element
        vmerge = tc.xpath(".//w:vMerge")
        return vmerge and vmerge[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "restart"

    # 工具方法：判断单元格是否独占整行
    def is_full_row_cell(row, target_cell):
        target_element = target_cell._element
        return all(cell._element is target_element for cell in row.cells)

    # 工具方法：找到当前单元格右边的"视觉上的一个格子"（处理合并单元格）
    def find_right_cell(cells, current_idx, current_cell):
        """
        找到当前单元格右边的"视觉上的一个格子"
        处理合并单元格的情况：跳过与当前单元格相同的合并单元格

        Args:
            cells: 行中的所有单元格列表
            current_idx: 当前单元格索引
            current_cell: 当前单元格对象

        Returns:
            右边的单元格对象，如果不存在则返回 None
        """
        current_element = current_cell._element

        # 从当前位置向右遍历，找到第一个不同的单元格
        for idx in range(current_idx + 1, len(cells)):
            if cells[idx]._element is not current_element:
                return cells[idx]

        return None

    # 工具方法：根据 mapping_data 中的 fillmapping 配置，手动指定单元格填充占位符
    def apply_manual_fill_mappings(doc, mapping_data, placeholders_dict):
        """
        根据 mapping_data 中的 fillmapping 配置，手动指定单元格填充占位符

        Args:
            doc: Document 对象
            mapping_data: 解析后的映射配置
            placeholders_dict: 占位符字典，用于收集所有占位符
        """
        if not mapping_data or 'fillmapping' not in mapping_data:
            return

        fillmapping = mapping_data.get('fillmapping', [])

        for mapping in fillmapping:
            title = mapping.get('title', '').strip()
            position = mapping.get('position', '').upper()
            mapfield = mapping.get('mapfield', '').strip()

            if not title or not mapfield:
                continue

            if position not in ('RIGHT', 'SELF'):
                continue  # 目前只支持 RIGHT 和 SELF

            # 在所有表格中查找包含 title 的单元格
            found = False
            for table in doc.tables:
                if found:
                    break
                for row in table.rows:
                    if found:
                        break
                    cells = row.cells
                    for cell_idx, cell in enumerate(cells):
                        if found:
                            break
                        if title in cell.text:
                            # 根据 position 类型确定目标单元格
                            if position == 'RIGHT':
                                target_cell = find_right_cell(cells, cell_idx, cell)
                            elif position == 'SELF':
                                target_cell = cell
                            else:
                                target_cell = None

                            if target_cell:
                                placeholder_value = f"{{{{{mapfield}}}}}"
                                update_cell_text_preserving_format(target_cell, placeholder_value)
                                placeholders_dict[mapfield] = ""
                                found = True

    # 工具方法：填充表格（右填充 + 下填充，检测合并单元格，独立单元格）
    def fill_table(matrix, table, above_paragraphs):
        rows = len(matrix)
        if rows == 0:
            print("Matrix is empty!")  # 调试日志
            return matrix, set()

        cols = len(matrix[0])
        right_filled_cells = set()
        down_filled_cells = set()
        print("Initial table matrix:")  # 打印整个表格矩阵

        # 第一阶段：右填充
        for r in range(rows):
            for c in range(1, cols):
                if not matrix[r][c].strip():
                    left_cell = table.rows[r].cells[c - 1]
                    current_cell = table.rows[r].cells[c]

                    # 如果左侧单元格为下合并的起始单元格，则跳过当前单元格的右填充
                    if is_start_of_down_merge(left_cell) or is_merged_cell(current_cell):
                        continue

                    if matrix[r][c - 1].strip():
                        left_value = wrap_placeholder(matrix[r][c - 1])
                        matrix[r][c] = left_value
                        right_filled_cells.add((r, c))

        # 第二阶段：下填充（连续传播填充逻辑）
        for c in range(cols):  # 遍历每一列
            for r in range(1, rows):  # 从第二行开始
                current_cell = table.rows[r].cells[c]
                top_cell = table.rows[r - 1].cells[c]

                # 如果当前单元格为空，且不是合并单元格，并且上方单元格有内容
                if not matrix[r][c].strip() and not is_merged_cell(current_cell) and matrix[r - 1][c].strip():
                    matrix[r][c] = wrap_placeholder(matrix[r - 1][c])
                    down_filled_cells.add((r, c))

        # 第三阶段：处理独立单元格
        for r in range(rows):
            for c in range(cols):
                if not matrix[r][c].strip() and is_isolated_cell(matrix, r, c):
                    if above_paragraphs:
                        description = above_paragraphs[-1]
                        placeholder_key = normalize_placeholder_key(description)
                        placeholder_value = wrap_placeholder(placeholder_key)
                        matrix[r][c] = placeholder_value
                        down_filled_cells.add((r, c))

        return matrix, down_filled_cells

    # 工具方法：按列为下填充的占位符添加下划线序号
    def add_sequential_suffixes_by_column(matrix, filled_cells):
        rows = len(matrix)
        if rows == 0:
            return matrix

        max_cols = max(len(row) for row in matrix)
        for c in range(max_cols):
            placeholders_seen = {}
            for r in range(rows):
                if (r, c) in filled_cells:
                    val = matrix[r][c].strip()
                    if is_placeholder_format(val):
                        raw_key = val.strip('{').strip('}')
                        if raw_key not in placeholders_seen:
                            placeholders_seen[raw_key] = 1
                            matrix[r][c] = f"{{{{{raw_key}_1}}}}"
                        else:
                            placeholders_seen[raw_key] += 1
                            suffix_index = placeholders_seen[raw_key]
                            matrix[r][c] = f"{{{{{raw_key}_{suffix_index}}}}}"
        return matrix

    def insert_placeholder_preserving_layout(cell, placeholder: str):
        """
        将占位符插入到单元格中，尽量复用现有段落并保持原有对齐和样式。
        逻辑：识别顶部内容段落数量，并在其后的第一个合适位置插入新的段落。
        """
        if not placeholder:
            return

        paragraphs = list(cell.paragraphs)
        if not paragraphs:
            new_para = cell.add_paragraph('')
            _replace_paragraph_text(new_para, placeholder)
            return

        paragraph_texts = [p.text for p in paragraphs]
        aggregated_text = "\n".join(paragraph_texts)
        if placeholder in aggregated_text:
            return

        top_content = extract_top_content(aggregated_text)

        if not top_content:
            target_paragraph = paragraphs[0].insert_paragraph_before('')
            head_style = paragraphs[0].style
            if head_style:
                target_paragraph.style = head_style
            if paragraphs[0].alignment is not None:
                target_paragraph.alignment = paragraphs[0].alignment
            else:
                target_paragraph.alignment = None
            _replace_paragraph_text(target_paragraph, placeholder)
            return

        lines = paragraph_texts
        top_lines = top_content.split('\n')
        insert_index = len(top_lines)

        if insert_index < len(lines) and lines[insert_index].strip() == "":
            insert_index += 1

        base_idx = max(0, min(len(paragraphs) - 1, len(top_lines) - 1))
        base_paragraph = paragraphs[base_idx]
        base_style = base_paragraph.style
        base_alignment = base_paragraph.alignment

        if insert_index >= len(paragraphs):
            target_paragraph = cell.add_paragraph('')
        else:
            target_paragraph = paragraphs[insert_index].insert_paragraph_before('')

        if base_style:
            target_paragraph.style = base_style
        if base_alignment is not None:
            target_paragraph.alignment = base_alignment
        else:
            target_paragraph.alignment = None

        _replace_paragraph_text(target_paragraph, placeholder)

    # Step 1: 读取文件（使用已读取的内容）
    input_doc = Document(BytesIO(input_file_content))

    # Step 2: 处理表格
    placeholders = {}
    all_tables_original = []
    all_tables_result = []
    all_down_filled_cells = []

    # 遍历输入文档中的所有表格，提取其内容到一个三维列表中
    for table in input_doc.tables:
        # 使用列表推导式，将表格的每个单元格的文本内容去空格后，按行存储到二维列表中
        original_matrix = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        # 将提取的表格内容添加到所有表格的原始内容列表中
        all_tables_original.append(original_matrix)

    # 遍历所有原始表格内容，进行处理
    for idx, table in enumerate(input_doc.tables):
        # 提取当前表格的原始矩阵
        original_matrix = all_tables_original[idx]

        # 深拷贝原始矩阵以避免修改原始数据
        result_matrix = [row[:] for row in original_matrix]

        above_paragraphs = get_table_above_paragraphs(input_doc, idx)

        # 调用 fill_table，传入 result_matrix 和当前的表格对象 table
        filled_matrix, down_filled_cells = fill_table(result_matrix, table, above_paragraphs)

        # 将填充后的矩阵添加到所有表格的处理结果列表中
        all_tables_result.append(filled_matrix)

        # 将向下填充的单元格列表添加到所有向下填充单元格的列表中
        all_down_filled_cells.append(down_filled_cells)

    # 遍历所有表格结果和向下填充的单元格，按列添加序列后缀
    for idx, (table_matrix, down_filled_cells) in enumerate(zip(all_tables_result, all_down_filled_cells)):
        # 使用add_sequential_suffixes_by_column函数按列添加序列后缀，更新表格结果
        all_tables_result[idx] = add_sequential_suffixes_by_column(table_matrix, down_filled_cells)

    # 遍历所有处理后的表格结果，进行占位符规范化处理
    for idx, table_matrix in enumerate(all_tables_result):
        # 使用 normalize_all_placeholders 函数规范化占位符，获取新矩阵和占位符映射
        new_matrix, key_map = normalize_all_placeholders(table_matrix)
        # 更新规范化后的矩阵
        all_tables_result[idx] = new_matrix

    # Step 3: 填充文档内容（只更新发生变化的单元格，避免丢失非文本内容）
    for t_idx, table in enumerate(input_doc.tables):
        result_matrix = all_tables_result[t_idx]
        original_matrix = all_tables_original[t_idx]
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # 只有当单元格内容发生变化时才更新，避免破坏原有的非文本内容（如图片、特殊格式等）
                original_text = original_matrix[r_idx][c_idx]
                new_text = result_matrix[r_idx][c_idx]
                if original_text != new_text:
                    update_cell_text_preserving_format(cell, new_text)

    # Step 3.5: 应用手动映射填充（根据 mapping_data 中的 fillmapping 配置）
    apply_manual_fill_mappings(input_doc, mapping_data, placeholders)

    # Step 4: 收集占位符 JSON 数据
    for result_matrix in all_tables_result:
        for row in result_matrix:
            for val in row:
                val = val.strip()
                if is_placeholder_format(val):
                    key = val.strip("{").strip("}")
                    placeholders[key] = ""

    # --- 以上是对空单元格的处理，总是以左侧、上方是否有单元格且单元格中有内容来判断要对当前空单元格要填什么内容 ---
    # --- 有些单元格中是有内容的，内容本身就提示了填充内容，现在要对有内容的单元格做处理 ---
    processed_cells = set()  # 用于跟踪已处理的合并单元格，避免重复处理

    for table_idx, table in enumerate(input_doc.tables):
        rows = table.rows
        for row_idx, row in enumerate(rows):
            cells = row.cells
            for cell_idx, cell in enumerate(cells):
                # 创建单元格的唯一标识
                cell_id = (table_idx, id(cell._element))

                # 如果这个单元格已经被处理过（可能是合并单元格），跳过
                if cell_id in processed_cells:
                    continue

                cell_text = cell.text

                # 仅处理独占整行的单元格
                if not is_full_row_cell(row, cell):
                    continue

                # 判断是否需要处理这个非空单元格
                if should_process_non_empty_cell(cell_text):
                    # 提取顶部内容
                    top_content = extract_top_content(cell_text)

                    if top_content:
                        # 使用新的完整键名生成函数
                        placeholder_key = generate_full_placeholder_key(top_content)

                        if placeholder_key:  # 确保键名不为空
                            # 生成占位符
                            placeholder_value = wrap_placeholder(placeholder_key)

                            # 在保持现有段落布局的同时插入占位符
                            insert_placeholder_preserving_layout(cell, placeholder_value)

                            # 将新生成的占位符添加到placeholders字典中
                            placeholders[placeholder_key] = ""

                            # 标记这个单元格已处理
                            processed_cells.add(cell_id)

                            # 如果是合并单元格，标记所有相关单元格为已处理
                            # 通过检查同一行中是否有相同的单元格对象来判断合并单元格
                            for other_cell_idx, other_cell in enumerate(cells):
                                if other_cell._element is cell._element and other_cell_idx != cell_idx:
                                    other_cell_id = (table_idx, id(other_cell._element))
                                    processed_cells.add(other_cell_id)

    # --- 处理段落中的下划线填充项 ---
    # 识别模式：标签文本 + 冒号 + 下划线区域（font.underline=True）
    # 例如："单   位：________________系（科、室）：________________"
    for paragraph in input_doc.paragraphs:
        process_paragraph_underline_placeholders(paragraph, placeholders)

    # Step 5: 返回处理后的文档和占位符
    output_stream = BytesIO()
    input_doc.save(output_stream)
    output_stream.seek(0)
    output_file_content = output_stream.getvalue()

    # 保存输出文件和占位符数据
    save_docx_log(request_id, "output", output_file_content, {"placeholders": placeholders})
    logger.info(f"[process_docx] 输出 placeholders: {json.dumps(placeholders, ensure_ascii=False)}")
    logger.info(f"[process_docx] 请求ID: {request_id} 处理完成")

    response_data = {
        "placeholders": placeholders,
        "file": base64.b64encode(output_file_content).decode("utf-8")
    }

    return get_json_result(retmsg="File processed successfully.", data=response_data)


@router.post("/fill_docx", summary="Word 文档填充接口", response_description="返回填充后的文档 Base64 数据")
async def fill_docx(
        file: UploadFile = File(...),
        data: str = Form(...)
):
    """
    ### POST `/v1/document/fill_docx` Word 文档填充接口

**功能描述**:
该接口接收用户上传的带占位符的 Word 文档（`.docx` 格式）和 JSON 数据，将 JSON 数据填充到文档中的占位符位置，并返回填充后的文档（Base64 编码）。

---

### 请求参数 (Request Parameters)

| 字段        | 类型          | 必填 | 描述                                    |
|-------------|---------------|------|-----------------------------------------|
| `file`      | `UploadFile`  | 是   | 用户上传的 Word 文档，格式必须为 `.docx` |
| `data`      | `string`      | 是   | JSON 格式的字符串，包含占位符对应的键值对 |

#### 示例请求 (Sample Request)
**表单数据**:
- 文件上传: `file` 上传一个包含占位符的 Word 文档。
- JSON 数据:
    ```json
    {
        "name": "张三",
        "date": "2024-12-11"
    }
    ```

---

### 响应 (Response)

#### 成功响应 (200)

- **响应格式**:
    - **`Content-Type: application/json`**
    - **示例**:
    ```json
    {
        "retmsg": "File filled successfully.",
        "data": {
            "file": "base64-encoded-filled-file-content"
        }
    }
    ```

- **字段说明**:
    | 字段       | 类型        | 描述                                     |
    |------------|-------------|------------------------------------------|
    | `retmsg`   | `string`    | 响应消息，表示处理结果。                 |
    | `data`     | `object`    | 包含填充后的文档的 Base64 编码数据。      |
    | `file`     | `string`    | Base64 编码的填充文档数据，需解码后保存为 `.docx` 文件 |

#### 错误响应

| 状态码 | 错误类型                     | 描述                                     |
|--------|------------------------------|------------------------------------------|
| 400    | `Invalid file format`        | 上传文件格式不是 `.docx`                 |
| 400    | `Invalid JSON data provided` | 提供的 `data` 不是有效的 JSON 格式       |
| 500    | `Internal server error`      | 填充过程中出现意外错误                   |

- **示例**:
    ```json
    {
        "detail": "Invalid file format. Only .docx files are supported."
    }
    ```

---

### 主要流程

1. **文件校验**:
    - 验证上传文件是否为 `.docx` 格式，如果格式错误，返回 400 错误。
2. **解析 JSON 数据**:
    - 使用 Python 内置的 `json.loads` 解析上传的 JSON 数据。
    - 如果 JSON 数据格式错误，返回 400 错误。
3. **填充文档**:
    - 使用 `python-docx-template` 将 JSON 数据渲染到文档中的占位符。
    - 占位符的格式需为 `{{key}}`。
4. **生成 Base64 编码**:
    - 将填充后的文档保存到内存流中，并转换为 Base64 编码字符串。
5. **返回响应**:
    - 返回包含 Base64 编码文档的 JSON 数据，客户端可解码后保存为 `.docx` 文件。

---

### 注意事项

- **占位符格式**:
    - 文档中的占位符需遵循 `{{key}}` 格式，`key` 必须与 JSON 数据中的键匹配。
- **文件格式限制**:
    - 仅支持 `.docx` 格式的 Word 文件，其他格式会返回错误。
- **Base64 文档解码**:
    - 响应返回的文档内容为 Base64 编码，客户端需解码后保存为 `.docx` 文件以查看填充结果。
- **异常处理**:
    - 如果填充过程中发生未预料的错误，将返回 500 错误，并包含具体错误描述。

---
    """
    # 生成请求唯一标识符
    request_id = f"fill_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
    logger.info(f"[fill_docx] 请求ID: {request_id}, 文件名: {file.filename}")

    # 验证文件类型
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Invalid file format. Only .docx files are supported.")

    try:
        # Step 1: 解析 JSON 数据
        fill_data = json.loads(data)
        logger.info(f"[fill_docx] 输入 fill_data: {json.dumps(fill_data, ensure_ascii=False)}")

        # Step 2: 读取上传的文件
        input_file_content = await file.read()

        # 保存输入文件和 JSON 数据
        save_docx_log(request_id, "input", input_file_content, fill_data, file.filename)
        logger.info("[fill_docx] 输入文件已保存")

        doc = Document(BytesIO(input_file_content))

        # Step 3: 遍历表格，替换整格占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if not cell_text:
                        continue

                    def replace_placeholder(match):
                        key = match.group(1).strip()
                        value = fill_data.get(key, "")
                        if value is None:
                            value = ""
                        return str(value)

                    replaced_text, replaced_count = PLACEHOLDER_PATTERN.subn(replace_placeholder, cell_text)
                    if replaced_count > 0:
                        update_cell_text_preserving_format(cell, replaced_text)

        # Step 4: 遍历段落，替换下划线占位符
        for paragraph in doc.paragraphs:
            fill_paragraph_underline_placeholders(paragraph, fill_data)

        # Step 5: 保存填充后的文档到内存流
        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        output_file_content = output_stream.getvalue()

        # 保存输出文件
        save_docx_log(request_id, "output", output_file_content)
        logger.info(f"[fill_docx] 请求ID: {request_id} 处理完成")

        # Step 6: 将文档转为 Base64 数据
        base64_encoded_file = base64.b64encode(output_file_content).decode("utf-8")

        # Step 7: 返回结果
        response_data = {
            "file": base64_encoded_file
        }
        return get_json_result(retmsg="File filled successfully.", data=response_data)

    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON data provided.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {e!s}")
