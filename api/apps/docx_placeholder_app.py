import base64
import json
import logging
import os
import re
import shutil
import uuid
from datetime import datetime
from io import BytesIO

from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, UploadFile

from api.utils.api_utils import get_json_result

router = APIRouter()
logger = logging.getLogger(__name__)

# 临时文件存放目录
DOCX_LOG_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "temp", "docx_placeholder_logs")
# 最大保留的请求数量
DOCX_LOG_MAX_REQUESTS = 10

# 占位符匹配模式：{{$数字}}
PLACEHOLDER_PATTERN = re.compile(r"\{\{\$(\d+)\}\}")


def ensure_log_dir():
    """确保日志目录存在"""
    if not os.path.exists(DOCX_LOG_DIR):
        os.makedirs(DOCX_LOG_DIR)


def cleanup_old_logs():
    """
    清理旧的日志目录，只保留最近 DOCX_LOG_MAX_REQUESTS 次请求的数据
    按目录创建时间排序，删除最旧的
    """
    if not os.path.exists(DOCX_LOG_DIR):
        return

    # 获取所有子目录
    subdirs = []
    for name in os.listdir(DOCX_LOG_DIR):
        path = os.path.join(DOCX_LOG_DIR, name)
        if os.path.isdir(path):
            # 获取目录创建时间
            ctime = os.path.getctime(path)
            subdirs.append((path, ctime))

    # 如果目录数量未超过限制，不需要清理
    if len(subdirs) <= DOCX_LOG_MAX_REQUESTS:
        return

    # 按创建时间排序（最旧的在前）
    subdirs.sort(key=lambda x: x[1])

    # 删除最旧的目录，保留最近 DOCX_LOG_MAX_REQUESTS 个
    dirs_to_delete = subdirs[:-DOCX_LOG_MAX_REQUESTS]
    for dir_path, _ in dirs_to_delete:
        try:
            shutil.rmtree(dir_path)
            logger.info(f"[cleanup] 已删除旧日志目录: {os.path.basename(dir_path)}")
        except Exception as e:
            logger.warning(f"[cleanup] 删除目录失败 {dir_path}: {e}")


def save_docx_log(request_id: str, stage: str, file_content: bytes, json_data: dict = None, filename: str = None):
    """
    保存 docx 处理日志

    Args:
        request_id: 请求唯一标识符
        stage: 阶段标识，如 "input", "output"
        file_content: 文件内容（bytes）
        json_data: JSON 数据（可选）
        filename: 原始文件名（可选）
    """
    ensure_log_dir()

    # 在保存新日志前清理旧日志
    if stage == "input":
        cleanup_old_logs()

    # 创建请求专属目录
    request_dir = os.path.join(DOCX_LOG_DIR, request_id)
    if not os.path.exists(request_dir):
        os.makedirs(request_dir)

    # 保存文件
    file_suffix = f"_{filename}" if filename else ""
    docx_path = os.path.join(request_dir, f"{stage}{file_suffix}.docx")
    with open(docx_path, "wb") as f:
        f.write(file_content)

    # 保存 JSON 数据（如果有）
    if json_data is not None:
        json_path = os.path.join(request_dir, f"{stage}_data.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)

    return request_dir


def analyze_document_structure(doc: Document) -> dict:
    """
    分析文档结构，生成调试信息

    Args:
        doc: python-docx Document 对象

    Returns:
        文档结构分析结果
    """
    structure = {
        "body_paragraphs": [],
        "tables": [],
        "summary": {
            "total_body_paragraphs": 0,
            "total_tables": 0,
            "total_table_cells": 0,
            "placeholders_found": [],
            "potentially_split_placeholders": [],
        },
    }

    all_placeholders = set()

    # 分析正文段落
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_info = _analyze_paragraph(paragraph, f"body_para{para_idx}")
        structure["body_paragraphs"].append(para_info)

        # 收集占位符
        if para_info["contains_placeholder_pattern"]:
            found = extract_placeholders_from_text(para_info["full_text"])
            all_placeholders.update(found)

        # 检查是否有拆分的占位符
        if para_info["is_split"]:
            structure["summary"]["potentially_split_placeholders"].append(
                {
                    "location": para_info["location"],
                    "full_text": para_info["full_text"],
                    "runs_detail": para_info["runs_detail"],
                }
            )

    # 分析表格
    for table_idx, table in enumerate(doc.tables):
        table_info = {"table_index": table_idx, "rows": []}

        for row_idx, row in enumerate(table.rows):
            row_info = {"row_index": row_idx, "cells": []}

            for cell_idx, cell in enumerate(row.cells):
                cell_info = {"cell_index": cell_idx, "paragraphs": []}

                for para_idx, paragraph in enumerate(cell.paragraphs):
                    location = f"table{table_idx}_row{row_idx}_cell{cell_idx}_para{para_idx}"
                    para_info = _analyze_paragraph(paragraph, location)
                    cell_info["paragraphs"].append(para_info)

                    # 收集占位符
                    if para_info["contains_placeholder_pattern"]:
                        found = extract_placeholders_from_text(para_info["full_text"])
                        all_placeholders.update(found)

                    # 检查是否有拆分的占位符
                    if para_info["is_split"]:
                        structure["summary"]["potentially_split_placeholders"].append(
                            {
                                "location": para_info["location"],
                                "full_text": para_info["full_text"],
                                "runs_detail": para_info["runs_detail"],
                            }
                        )

                row_info["cells"].append(cell_info)
                structure["summary"]["total_table_cells"] += 1

            table_info["rows"].append(row_info)

        structure["tables"].append(table_info)

    # 更新汇总
    structure["summary"]["total_body_paragraphs"] = len(doc.paragraphs)
    structure["summary"]["total_tables"] = len(doc.tables)
    structure["summary"]["placeholders_found"] = [f"${num}" for num in sorted(all_placeholders)]
    structure["summary"]["split_placeholder_count"] = len(structure["summary"]["potentially_split_placeholders"])

    return structure


def _analyze_paragraph(paragraph, location: str) -> dict:
    """
    分析单个段落的结构

    Args:
        paragraph: python-docx Paragraph 对象
        location: 段落位置标识

    Returns:
        段落分析结果
    """
    para_text = paragraph.text if paragraph.text else ""

    runs_detail = []
    for i, run in enumerate(paragraph.runs):
        run_info = {
            "run_index": i,
            "text": run.text if run.text else "",
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline is not None and run.underline != False,
        }
        runs_detail.append(run_info)

    contains_placeholder = bool(PLACEHOLDER_PATTERN.search(para_text)) if para_text else False

    # 检查占位符是否被拆分
    is_split = False
    if contains_placeholder:
        # 检查是否有任何单个 run 包含完整的占位符
        has_complete_in_run = any(PLACEHOLDER_PATTERN.search(run["text"]) for run in runs_detail)
        if not has_complete_in_run:
            is_split = True

    return {
        "location": location,
        "full_text": para_text,
        "runs_count": len(paragraph.runs),
        "runs_detail": runs_detail,
        "contains_placeholder_pattern": contains_placeholder,
        "is_split": is_split,
    }


def save_document_structure(request_id: str, structure: dict):
    """
    保存文档结构分析结果到文件

    Args:
        request_id: 请求唯一标识符
        structure: 文档结构分析结果
    """
    request_dir = os.path.join(DOCX_LOG_DIR, request_id)
    if not os.path.exists(request_dir):
        os.makedirs(request_dir)

    structure_path = os.path.join(request_dir, "document_structure.json")
    with open(structure_path, "w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)

    # 输出汇总日志
    summary = structure["summary"]
    logger.info("[document_structure] ===== 文档结构分析 =====")
    logger.info(f"[document_structure] 正文段落数: {summary['total_body_paragraphs']}")
    logger.info(f"[document_structure] 表格数: {summary['total_tables']}")
    logger.info(f"[document_structure] 表格单元格数: {summary['total_table_cells']}")
    logger.info(f"[document_structure] 发现占位符: {summary['placeholders_found']}")
    logger.info(f"[document_structure] 疑似被拆分的占位符数: {summary['split_placeholder_count']}")

    if summary["potentially_split_placeholders"]:
        logger.warning("[document_structure] ⚠️ 发现可能被拆分的占位符:")
        for item in summary["potentially_split_placeholders"]:
            logger.warning(f"[document_structure]   位置: {item['location']}")
            logger.warning(f"[document_structure]   完整文本: {item['full_text']}")
            logger.warning("[document_structure]   Run 拆分情况:")
            for run in item["runs_detail"]:
                logger.warning(f"[document_structure]     Run[{run['run_index']}]: '{run['text']}'")

    logger.info(f"[document_structure] 结构分析已保存至: {structure_path}")

    return structure_path


def extract_placeholders_from_text(text: str) -> set:
    """
    从文本中提取占位符编号

    Args:
        text: 待检测的文本

    Returns:
        占位符编号集合（例如：{1, 2, 3}）
    """
    matches = PLACEHOLDER_PATTERN.findall(text)
    return {int(num) for num in matches}


def extract_placeholders_from_document(doc: Document) -> list:
    """
    从 Word 文档中提取所有占位符

    Args:
        doc: python-docx Document 对象

    Returns:
        按数字顺序排序的占位符列表
        格式：[{"pos": "$1", "alias": ""}, {"pos": "$2", "alias": ""}]
    """
    placeholder_numbers = set()

    # 从段落中提取
    for paragraph in doc.paragraphs:
        if paragraph.text:
            found = extract_placeholders_from_text(paragraph.text)
            placeholder_numbers.update(found)

    # 从表格单元格中提取
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    found = extract_placeholders_from_text(cell.text)
                    placeholder_numbers.update(found)

    # 按数字排序并格式化返回
    sorted_numbers = sorted(placeholder_numbers)
    result = [{"pos": f"${num}", "alias": ""} for num in sorted_numbers]

    return result


def fill_placeholders_in_paragraph(paragraph, fill_dict: dict) -> bool:
    """
    填充段落中的占位符，保留原有格式（如下划线）
    支持处理占位符被拆分到多个 run 的情况

    Args:
        paragraph: python-docx Paragraph 对象
        fill_dict: 填充数据字典，key 为 "$1" 格式，value 为填充内容

    Returns:
        是否进行了替换
    """
    para_text = paragraph.text if paragraph.text else ""

    if not para_text:
        return False

    # 检查是否包含占位符
    if not PLACEHOLDER_PATTERN.search(para_text):
        return False

    # 构建 run 的位置映射
    runs = paragraph.runs
    run_texts = [run.text if run.text else "" for run in runs]

    # 合并所有 run 的文本
    combined_text = "".join(run_texts)

    # 检查合并后的文本是否包含占位符
    if not PLACEHOLDER_PATTERN.search(combined_text):
        return False

    # 构建字符到 run 的映射
    char_to_run = []  # char_to_run[i] = (run_index, position_in_run)
    for run_idx, text in enumerate(run_texts):
        for char_idx in range(len(text)):
            char_to_run.append((run_idx, char_idx))

    # 找到所有占位符的位置
    matches = list(PLACEHOLDER_PATTERN.finditer(combined_text))

    if not matches:
        return False

    replaced = False

    # 从后往前替换，避免位置偏移问题
    for match in reversed(matches):
        start_pos = match.start()
        end_pos = match.end()
        num = match.group(1)
        key = f"${num}"
        value = fill_dict.get(key, "")
        if value is None:
            value = ""

        # 找出占位符跨越了哪些 run
        start_run_idx, start_char_idx = char_to_run[start_pos]
        end_run_idx, end_char_idx = char_to_run[end_pos - 1]

        if start_run_idx == end_run_idx:
            # 占位符在单个 run 内，直接替换
            run = runs[start_run_idx]
            original_text = run.text
            new_text = original_text[:start_char_idx] + str(value) + original_text[end_char_idx + 1 :]
            run.text = new_text
        else:
            # 占位符跨越多个 run
            # 策略：在第一个 run 中放置替换值，清空中间和最后的 run 中属于占位符的部分

            # 处理第一个 run：保留占位符之前的内容 + 替换值
            first_run = runs[start_run_idx]
            first_run.text = first_run.text[:start_char_idx] + str(value)

            # 处理中间的 run（如果有）：清空
            for mid_run_idx in range(start_run_idx + 1, end_run_idx):
                runs[mid_run_idx].text = ""

            # 处理最后一个 run：保留占位符之后的内容
            last_run = runs[end_run_idx]
            last_run.text = last_run.text[end_char_idx + 1 :]

        replaced = True

    return replaced


def fill_placeholders_in_cell(cell, fill_dict: dict) -> bool:
    """
    填充表格单元格中的占位符

    Args:
        cell: python-docx Cell 对象
        fill_dict: 填充数据字典

    Returns:
        是否进行了替换
    """
    if not cell.text:
        return False

    replaced = False

    # 遍历单元格中的所有段落
    for paragraph in cell.paragraphs:
        if fill_placeholders_in_paragraph(paragraph, fill_dict):
            replaced = True

    return replaced


def fill_placeholders_in_document(doc: Document, fill_data: list) -> int:
    """
    填充文档中的所有占位符

    Args:
        doc: python-docx Document 对象
        fill_data: 填充数据列表
                   格式：[{"pos": "$1", "alias": "xxx", "fill": "张三"}, ...]

    Returns:
        替换的占位符数量
    """
    # 将填充数据转换为字典，方便查找
    fill_dict = {}
    for item in fill_data:
        pos = item.get("pos", "")
        fill = item.get("fill", "")
        if pos:
            fill_dict[pos] = fill

    replace_count = 0

    # 填充段落
    for paragraph in doc.paragraphs:
        if fill_placeholders_in_paragraph(paragraph, fill_dict):
            replace_count += 1

    # 填充表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if fill_placeholders_in_cell(cell, fill_dict):
                    replace_count += 1

    return replace_count


@router.post("/extract_placeholders", summary="提取 Word 文档中的占位符", response_description="返回占位符列表")
async def extract_placeholders(file: UploadFile = File(...)):
    """
    ### POST `/v1/docx_placeholder/extract_placeholders` 提取 Word 文档占位符接口

    **功能描述**:
    该接口接收用户上传的 Word 文件，提取文档中所有的占位符（格式为 {{$1}}, {{$2}} 等），
    并返回按数字顺序排序的占位符列表。

    ---

    ### 请求体 (Request Body)

    | 字段   | 类型         | 必填 | 描述                                    |
    |--------|--------------|------|-----------------------------------------|
    | `file` | `UploadFile` | 是   | 用户上传的 Word 文档，格式必须为 `.docx` |

    ---

    ### 响应 (Response)

    #### 成功响应 (200)

    - **响应格式**:
        - **`Content-Type: application/json`**
        - **示例**:
        ```json
        {
            "retcode": 0,
            "retmsg": "Placeholders extracted successfully.",
            "data": [
                {"pos": "$1", "alias": ""},
                {"pos": "$2", "alias": ""}
            ]
        }
        ```

    - **字段说明**:
        | 字段   | 类型   | 描述                                     |
        |--------|--------|------------------------------------------|
        | `pos`  | `string` | 占位符位置标识，格式为 "$数字"           |
        | `alias`| `string` | 占位符别名，当前为空字符串               |

    ---

    #### 错误响应

    | 状态码 | 错误类型              | 描述                                |
    |--------|-----------------------|-------------------------------------|
    | 400    | `Invalid file format` | 当上传的文件格式不是 `.docx` 时返回 |

    ---

    ### 主要流程

    1. **文件校验**: 验证上传文件的格式是否为 `.docx`
    2. **读取文档**: 使用 `python-docx` 读取文档内容
    3. **提取占位符**: 从段落和表格单元格中提取所有 `{{$数字}}` 格式的占位符
    4. **排序返回**: 按数字顺序排序后返回 JSON 数组

    ---

    ### 注意事项

    - **占位符格式**: 必须严格遵循 `{{$数字}}` 格式，例如 `{{$1}}`, `{{$2}}`
    - **支持位置**: 段落文本、表格单元格
    - **不支持位置**: 页眉、页脚（暂不支持）
    - **去重**: 相同的占位符只会在结果中出现一次
    """
    # 生成请求唯一标识符
    request_id = f"extract_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
    logger.info(f"[extract_placeholders] 请求ID: {request_id}, 文件名: {file.filename}")

    # 验证文件名和格式
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Invalid file format. Only .docx files are supported.")

    # 读取上传的文件内容
    input_file_content = await file.read()

    # 保存输入文件
    save_docx_log(request_id, "input", input_file_content, filename=file.filename)
    logger.info("[extract_placeholders] 输入文件已保存")

    # 读取文档
    doc = Document(BytesIO(input_file_content))

    # 提取占位符
    placeholders = extract_placeholders_from_document(doc)

    # 保存提取结果
    save_docx_log(request_id, "output", input_file_content, {"placeholders": placeholders})
    logger.info(f"[extract_placeholders] 提取到 {len(placeholders)} 个占位符: {json.dumps(placeholders, ensure_ascii=False)}")
    logger.info(f"[extract_placeholders] 请求ID: {request_id} 处理完成")

    return get_json_result(retmsg="Placeholders extracted successfully.", data=placeholders)


@router.post("/fill_placeholders", summary="填充 Word 文档中的占位符", response_description="返回填充后的文档 Base64 数据")
async def fill_placeholders(file: UploadFile = File(...), data: str = Form(...)):
    """
    ### POST `/v1/docx_placeholder/fill_placeholders` 填充 Word 文档占位符接口

    **功能描述**:
    该接口接收用户上传的带占位符的 Word 文档和填充数据，将占位符替换为对应的内容，
    并返回填充后的文档（Base64 编码）。

    ---

    ### 请求参数 (Request Parameters)

    | 字段   | 类型         | 必填 | 描述                                    |
    |--------|--------------|------|-----------------------------------------|
    | `file` | `UploadFile` | 是   | 用户上传的 Word 文档，格式必须为 `.docx` |
    | `data` | `string`     | 是   | JSON 格式的字符串，包含占位符填充数据   |

    #### 示例请求 (Sample Request)
    **表单数据**:
    - 文件上传: `file` 上传一个包含占位符的 Word 文档
    - JSON 数据:
        ```json
        [
            {"pos": "$1", "alias": "姓名", "fill": "张三"},
            {"pos": "$2", "alias": "日期", "fill": "2024-12-17"}
        ]
        ```

    ---

    ### 响应 (Response)

    #### 成功响应 (200)

    - **响应格式**:
        - **`Content-Type: application/json`**
        - **示例**:
        ```json
        {
            "retcode": 0,
            "retmsg": "File filled successfully.",
            "data": {
                "file": "base64-encoded-filled-file-content"
            }
        }
        ```

    - **字段说明**:
        | 字段   | 类型     | 描述                                           |
        |--------|----------|------------------------------------------------|
        | `file` | `string` | Base64 编码的填充后文档，需解码后保存为 `.docx` |

    #### 错误响应

    | 状态码 | 错误类型                     | 描述                                     |
    |--------|------------------------------|------------------------------------------|
    | 400    | `Invalid file format`        | 上传文件格式不是 `.docx`                 |
    | 400    | `Invalid JSON data provided` | 提供的 `data` 不是有效的 JSON 格式       |
    | 500    | `Internal server error`      | 填充过程中出现意外错误                   |

    ---

    ### 主要流程

    1. **文件校验**: 验证上传文件是否为 `.docx` 格式
    2. **解析 JSON 数据**: 解析填充数据
    3. **填充文档**: 遍历段落和表格，替换所有匹配的占位符
    4. **生成 Base64 编码**: 将填充后的文档转换为 Base64 编码
    5. **返回响应**: 返回包含 Base64 编码文档的 JSON 数据

    ---

    ### 注意事项

    - **占位符格式**: 文档中的占位符必须为 `{{$数字}}` 格式
    - **未提供的占位符**: 如果文档中的占位符在 data 中没有对应的 fill 值，将替换为空字符串
    - **不存在的占位符**: 如果 data 中提供的占位符在文档中不存在，将被忽略
    - **Base64 解码**: 客户端需要将返回的 Base64 字符串解码后保存为 `.docx` 文件
    """
    # 生成请求唯一标识符
    request_id = f"fill_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
    logger.info(f"[fill_placeholders] 请求ID: {request_id}, 文件名: {file.filename}")

    # 验证文件类型
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Invalid file format. Only .docx files are supported.")

    try:
        # Step 1: 解析 JSON 数据
        fill_data = json.loads(data)
        logger.info(f"[fill_placeholders] 输入 fill_data: {json.dumps(fill_data, ensure_ascii=False)}")

        # Step 2: 读取上传的文件
        input_file_content = await file.read()

        # 保存输入文件和 JSON 数据
        save_docx_log(request_id, "input", input_file_content, fill_data, file.filename)
        logger.info("[fill_placeholders] 输入文件已保存")

        # Step 3: 读取文档
        doc = Document(BytesIO(input_file_content))

        # Step 4: 分析文档结构（独立的调试流程）
        doc_structure = analyze_document_structure(doc)
        save_document_structure(request_id, doc_structure)

        # Step 5: 填充占位符
        replace_count = fill_placeholders_in_document(doc, fill_data)
        logger.info(f"[fill_placeholders] 共替换了 {replace_count} 处占位符")

        # Step 6: 保存填充后的文档到内存流
        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        output_file_content = output_stream.getvalue()

        # 保存输出文件
        save_docx_log(request_id, "output", output_file_content)
        logger.info(f"[fill_placeholders] 请求ID: {request_id} 处理完成")

        # Step 7: 将文档转为 Base64 数据
        base64_encoded_file = base64.b64encode(output_file_content).decode("utf-8")

        # Step 8: 返回结果
        response_data = {"file": base64_encoded_file}
        return get_json_result(retmsg="File filled successfully.", data=response_data)

    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON data provided.")
    except Exception as e:
        logger.error(f"[fill_placeholders] 处理失败: {e!s}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An error occurred: {e!s}")
