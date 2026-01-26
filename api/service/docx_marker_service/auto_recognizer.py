# coding=utf-8
"""自动识别文档中的待填充位置"""

import re
from docx import Document
from docx.oxml.ns import qn

from .models import Placeholder, TableColumn, TableConfig
from .parser import parse_table


# 冒号模式（中英文冒号）
COLON_PATTERN = re.compile(r'[：:]')
HEADING_PATTERN = re.compile(r'^\s*\d+[\.、．]')
ROW_INDEX_PATTERN = re.compile(r'/row\[(\d+)]')


def normalize_placeholder_key(key: str) -> str:
    """标准化占位符键名，保留中文字符、数字、字母、下划线和常用符号"""
    # 移除多余空格
    normalized = re.sub(r'\s+', '', key)
    # 保留中文、数字、字母、下划线和常用符号（括号、顿号等）
    normalized = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9_（）()、]', '', normalized)
    return normalized


def extract_underline_placeholders_from_paragraph(paragraph):
    """
    从段落中提取带下划线的填充项
    """
    placeholders = []
    runs = list(paragraph.runs)

    if not runs:
        return placeholders

    # 构建run信息列表
    run_info = []
    char_offset = 0
    for i, run in enumerate(runs):
        text = run.text or ""
        is_underline = bool(run.font.underline)
        run_info.append({
            'index': i,
            'text': text,
            'is_underline': is_underline,
            'char_start': char_offset,
            'char_end': char_offset + len(text),
            'run': run
        })
        char_offset += len(text)

    # 找到所有下划线区域
    underline_regions = []
    i = 0
    while i < len(run_info):
        if run_info[i]['is_underline']:
            region_start = i
            region_end = i
            while region_end + 1 < len(run_info) and run_info[region_end + 1]['is_underline']:
                region_end += 1
            underline_regions.append({
                'start_index': region_start,
                'end_index': region_end,
                'run_indices': list(range(region_start, region_end + 1))
            })
            i = region_end + 1
        else:
            i += 1

    for region in underline_regions:
        label = _find_label_before_underline(run_info, region['start_index'])
        if label:
            placeholders.append({
                'label': label['text'],
                'underline_run_indices': region['run_indices'],
                'colon_run_index': label['colon_run_index'],
                'colon_position': label['colon_position']
            })

    return placeholders


def _find_label_before_underline(run_info, underline_start_index):
    text_before = ""
    run_text_mapping = []

    for i in range(underline_start_index):
        text = run_info[i]['text']
        start_pos = len(text_before)
        text_before += text
        run_text_mapping.append((i, start_pos, start_pos + len(text)))

    if not text_before:
        return None

    colon_match = None
    for match in COLON_PATTERN.finditer(text_before):
        colon_match = match

    if colon_match is None:
        return None

    colon_pos = colon_match.start()
    colon_run_index = None
    for run_idx, start, end in run_text_mapping:
        if start <= colon_pos < end:
            colon_run_index = run_idx
            break

    if colon_run_index is None:
        return None

    label_start = 0
    for i in range(underline_start_index - 1, -1, -1):
        if run_info[i]['is_underline']:
            label_start = run_info[i]['char_end']
            break

    label_text = text_before[label_start:colon_pos].strip()
    if not label_text:
        return None

    return {
        'text': label_text,
        'colon_run_index': colon_run_index,
        'colon_position': colon_pos - run_text_mapping[colon_run_index][1]
    }


def get_cell_text(cell) -> str:
    parts = []
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text:
                parts.append(run.text)
    return ''.join(parts).strip()


def is_cell_empty(cell) -> bool:
    return get_cell_text(cell) == ""


def extract_row_index(path: str | None) -> int | None:
    if not path:
        return None
    match = ROW_INDEX_PATTERN.search(path)
    if match:
        return int(match.group(1))
    return None


def is_vertical_merge_continuation(cell, row_idx: int) -> bool:
    if cell.is_merged_origin:
        return False
    origin_row = extract_row_index(cell.merge_origin_path)
    return origin_row is not None and origin_row < row_idx


def is_data_row(row, row_idx: int) -> bool:
    for cell in row.cells:
        text = get_cell_text(cell)
        if not text:
            continue
        if is_vertical_merge_continuation(cell, row_idx):
            continue
        return False
    return True


def collect_data_row_groups(rows) -> list[tuple[int, int]]:
    groups = []
    start = None
    for idx, row in enumerate(rows):
        if is_data_row(row, idx):
            if start is None:
                start = idx
        else:
            if start is not None:
                groups.append((start, idx - 1))
                start = None
    if start is not None:
        groups.append((start, len(rows) - 1))
    return groups


def is_heading_text(text: str) -> bool:
    return bool(HEADING_PATTERN.match(text))


def detect_category_columns(table, start_row: int, end_row: int) -> list[int]:
    category_columns = []
    rows = table.rows

    if start_row >= len(rows):
        return category_columns

    first_row = rows[start_row]
    num_cells = len(first_row.cells)

    for col_idx in range(num_cells):
        cell = first_row.cells[col_idx]

        if cell.is_merged_origin and cell.row_span > 1:
            category_columns.append(col_idx)
            continue

        if (not cell.is_merged_origin) and cell.merge_origin_path:
            origin_row = extract_row_index(cell.merge_origin_path)
            if origin_row is not None and origin_row != start_row:
                category_columns.append(col_idx)
                continue
            continue

        if end_row <= start_row:
            continue

        cell_text = get_cell_text(cell)
        if not cell_text:
            continue

        if col_idx > 0:
            left_cell = first_row.cells[col_idx - 1]
            if get_cell_text(left_cell) == cell_text:
                continue

        is_repeating = True
        for row_idx in range(start_row + 1, min(end_row + 1, len(rows))):
            row_cell = rows[row_idx].cells[col_idx]
            if get_cell_text(row_cell) != cell_text:
                is_repeating = False
                break

        if is_repeating:
            category_columns.append(col_idx)

    return category_columns


def find_category_boundary(table, clicked_row_index: int, category_columns: list[int]) -> tuple[int, int]:
    rows = table.rows
    region_start = clicked_row_index
    region_end = clicked_row_index

    if not category_columns:
        for i in range(clicked_row_index - 1, -1, -1):
            row = rows[i]
            has_content = any(not is_cell_empty(cell) for cell in row.cells)
            if has_content:
                region_start = i
                break

        for i in range(clicked_row_index + 1, len(rows)):
            row = rows[i]
            is_empty = all(is_cell_empty(cell) for cell in row.cells)
            if is_empty:
                region_end = i
            else:
                break

        return region_start, region_end

    primary_col = category_columns[0]

    for i in range(clicked_row_index - 1, -1, -1):
        cell = rows[i].cells[primary_col]

        if (not cell.is_merged_origin) and cell.merge_origin_path:
            region_start = i
            continue

        if cell.is_merged_origin and cell.row_span > 1:
            if i + cell.row_span > clicked_row_index:
                region_start = i
            break

        clicked_text = get_cell_text(rows[clicked_row_index].cells[primary_col])
        cell_text = get_cell_text(cell)
        if cell_text and cell_text == clicked_text:
            region_start = i
        else:
            break

    start_cell = rows[region_start].cells[primary_col]
    if start_cell.is_merged_origin and start_cell.row_span > 1:
        region_end = region_start + start_cell.row_span - 1
    else:
        start_text = get_cell_text(start_cell)
        for i in range(clicked_row_index + 1, len(rows)):
            cell = rows[i].cells[primary_col]
            cell_text = get_cell_text(cell)
            if start_text and cell_text == start_text:
                region_end = i
            else:
                break

    return region_start, region_end


def is_title_row(row) -> bool:
    non_empty_cells = [cell for cell in row.cells if not is_cell_empty(cell)]
    if len(non_empty_cells) < 2:
        return False

    first_text = get_cell_text(non_empty_cells[0])
    if not first_text:
        return False

    return all(get_cell_text(cell) == first_text for cell in non_empty_cells)


def get_first_non_empty_text(row) -> str:
    for cell in row.cells:
        text = get_cell_text(cell)
        if text:
            return text
    return ""


def table_has_text(table) -> bool:
    for row in table.rows:
        for cell in row.cells:
            if get_cell_text(cell):
                return True
    return False


def get_full_width_origin_cell(row, num_cols: int):
    origin_cells = []
    for idx, cell in enumerate(row.cells):
        if cell.is_merged_origin and cell.col_span == num_cols:
            origin_cells.append((idx, cell))
    if len(origin_cells) != 1:
        return None
    return origin_cells[0]


def build_summary_placeholders(table, skip_rows: set[int]) -> tuple[list[Placeholder], set[int]]:
    placeholders = []
    rows = table.rows
    if not rows:
        return placeholders, skip_rows

    num_cols = len(rows[0].cells)
    for row_idx in range(len(rows) - 1):
        if row_idx in skip_rows or (row_idx + 1) in skip_rows:
            continue
        label_cell_info = get_full_width_origin_cell(rows[row_idx], num_cols)
        if not label_cell_info:
            continue
        _, label_cell = label_cell_info
        label_text = get_cell_text(label_cell)
        if not label_text:
            continue

        target_cell_info = get_full_width_origin_cell(rows[row_idx + 1], num_cols)
        if not target_cell_info:
            continue
        _, target_cell = target_cell_info
        if get_cell_text(target_cell):
            continue

        placeholders.append(Placeholder(
            path=target_cell.path,
            label=label_text,
            field_key="",
            type="summary",
        ))

        skip_rows.add(row_idx)
        row_span = target_cell.row_span if target_cell.is_merged_origin else 1
        for skip_row in range(row_idx + 1, min(row_idx + 1 + row_span, len(rows))):
            skip_rows.add(skip_row)

    return placeholders, skip_rows


def detect_table_region(table, clicked_row_index: int) -> dict | None:
    rows = table.rows
    if not rows:
        return None

    preliminary_data_end = clicked_row_index
    for i in range(clicked_row_index + 1, len(rows)):
        row = rows[i]
        if all(is_cell_empty(cell) for cell in row.cells[1:]):
            preliminary_data_end = i
        else:
            break

    category_columns = detect_category_columns(table, clicked_row_index, preliminary_data_end)
    region_start, region_end = find_category_boundary(table, clicked_row_index, category_columns)

    header_row = region_start
    title_row = None
    if header_row < len(rows) and is_title_row(rows[header_row]):
        title_row = header_row
        header_row += 1
    elif header_row > 0 and is_title_row(rows[header_row - 1]):
        title_row = header_row - 1

    data_start_row = max(clicked_row_index, header_row + 1)

    data_end_row = data_start_row
    for i in range(data_start_row, region_end + 1):
        row = rows[i]
        is_empty_data_row = all(
            (idx in category_columns) or is_cell_empty(cell)
            for idx, cell in enumerate(row.cells)
        )
        if is_empty_data_row:
            data_end_row = i
        elif i > data_start_row:
            break

    header_row_data = rows[header_row] if header_row < len(rows) else None
    columns: list[TableColumn] = []
    used_fallback_columns = False
    if header_row_data:
        for idx, cell in enumerate(header_row_data.cells):
            if idx in category_columns:
                continue
            if not cell.is_merged_origin:
                continue
            text = get_cell_text(cell)
            if text:
                columns.append(TableColumn(cell_index=idx, name=text))

    if not columns:
        used_fallback_columns = True
        fallback_row = rows[data_start_row] if data_start_row < len(rows) else rows[clicked_row_index]
        for idx, cell in enumerate(fallback_row.cells):
            if idx in category_columns:
                continue
            if cell.is_merged_origin:
                columns.append(TableColumn(cell_index=idx, name=f"列{idx + 1}"))

    return {
        "region_start": region_start,
        "header_row": header_row,
        "data_start_row": data_start_row,
        "data_end_row": data_end_row,
        "columns": columns,
        "category_columns": category_columns,
        "title_row": title_row,
        "region_end": region_end,
        "used_fallback_columns": used_fallback_columns,
    }


def build_table_placeholders(table, last_heading: str, last_paragraph: str) -> tuple[list[Placeholder], set[int]]:
    placeholders = []
    skip_rows: set[int] = set()
    if not table_has_text(table):
        return placeholders, skip_rows
    rows = table.rows
    data_groups = collect_data_row_groups(rows)
    if not data_groups:
        return placeholders, skip_rows

    last_region_end = -1
    for group_start, _group_end in data_groups:
        if group_start <= last_region_end:
            continue

        region = detect_table_region(table, group_start)
        if not region:
            continue

        last_region_end = max(last_region_end, region["region_end"])
        if region["data_end_row"] < region["data_start_row"]:
            continue
        if not region["columns"]:
            continue
        data_row_count = region["data_end_row"] - region["data_start_row"] + 1
        if len(region["columns"]) < 2:
            if region["used_fallback_columns"] or data_row_count < 2:
                continue

        title_text = ""
        if region["title_row"] is not None:
            title_text = get_first_non_empty_text(rows[region["title_row"]])

        category_text = ""
        if region["category_columns"]:
            col_idx = region["category_columns"][0]
            category_text = get_cell_text(rows[region["header_row"]].cells[col_idx])

        label = title_text or category_text or last_heading
        if not label:
            header_row = rows[region["header_row"]]
            header_texts = [get_cell_text(cell) for cell in header_row.cells if get_cell_text(cell)]
            label = " / ".join(dict.fromkeys(header_texts))
        if not label:
            label = last_paragraph or "表格"

        anchor_col = region["columns"][0].cell_index if region["columns"] else (
            region["category_columns"][0] if region["category_columns"] else 0
        )
        anchor_path = rows[region["header_row"]].cells[anchor_col].path

        table_config = TableConfig(
            dynamic=True,
            header_row=region["header_row"],
            data_start_row=region["data_start_row"],
            data_end_row=region["data_end_row"],
            columns=region["columns"],
            category_columns=region["category_columns"],
        )
        placeholders.append(Placeholder(
            path=anchor_path,
            label=label,
            field_key="",
            type="dynamic_table",
            table_config=table_config,
        ))
        for row_idx in range(region["region_start"], region["region_end"] + 1):
            skip_rows.add(row_idx)

    return placeholders, skip_rows


def is_independent_cell(row, cell_idx, cell) -> bool:
    cells = row.cells
    current_element = cell._element
    if cell_idx > 0:
        left_cell = cells[cell_idx - 1]
        if left_cell._element is current_element:
            return False
    return True

def is_merged_cell(cell) -> bool:
    tc = cell._element
    vmerge = tc.xpath(".//w:vMerge")
    if not vmerge:
        return False
    return vmerge[0].get(qn("w:val")) != "restart"

def is_start_of_down_merge(cell) -> bool:
    tc = cell._element
    vmerge = tc.xpath(".//w:vMerge")
    if not vmerge:
        return False
    return vmerge[0].get(qn("w:val")) == "restart"

def auto_recognize_placeholders(doc_path: str) -> list[Placeholder]:
    doc = Document(doc_path)
    placeholders = []

    body = doc.element.body
    element_index = 0
    last_heading_text = ""
    last_paragraph_text = ""

    def process_element(child, index):
        nonlocal last_heading_text, last_paragraph_text
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            para = None
            for p in doc.paragraphs:
                if p._element == child:
                    para = p
                    break

            if para:
                underline_items = extract_underline_placeholders_from_paragraph(para)
                for item in underline_items:
                    label = item['label']
                    underline_run_indices = item['underline_run_indices']
                    if underline_run_indices:
                        first_run_idx = underline_run_indices[0]
                        path = f"body[{index}]/run[{first_run_idx}]"
                        placeholders.append(Placeholder(
                            path=path,
                            label=label,
                            field_key=""
                        ))
                para_text = para.text.strip()
                if para_text:
                    last_paragraph_text = para_text
                    if is_heading_text(para_text):
                        last_heading_text = para_text
            return True

        elif tag == 'tbl':
            table = None
            for tbl in doc.tables:
                if tbl._element == child:
                    table = tbl
                    break

            if table:
                table_path = f"body[{index}]"
                parsed_table = parse_table(table, table_path)
                table_placeholders, skip_rows = build_table_placeholders(
                    parsed_table,
                    last_heading_text,
                    last_paragraph_text
                )
                if table_placeholders:
                    placeholders.extend(table_placeholders)

                summary_placeholders, skip_rows = build_summary_placeholders(parsed_table, skip_rows)
                if summary_placeholders:
                    placeholders.extend(summary_placeholders)

                rows = table.rows
                num_rows = len(rows)
                num_cols = len(rows[0].cells) if num_rows > 0 else 0
                matrix = [[cell.text.strip() for cell in row.cells] for row in rows]

                # 右填充
                for r in range(num_rows):
                    if r in skip_rows:
                        continue
                    for c in range(1, num_cols):
                        if not matrix[r][c].strip():
                            current_cell = rows[r].cells[c]
                            if is_merged_cell(current_cell): continue
                            if not is_independent_cell(rows[r], c, current_cell): continue

                            source_value = None
                            for left_c in range(c - 1, -1, -1):
                                left_cell = rows[r].cells[left_c]
                                if is_start_of_down_merge(left_cell): break
                                if is_independent_cell(rows[r], left_c, left_cell) and matrix[r][left_c].strip():
                                    source_value = matrix[r][left_c]
                                    break

                            if source_value:
                                path = f"{table_path}/row[{r}]/cell[{c}]"
                                placeholders.append(Placeholder(path=path, label=source_value, field_key=""))

                # 收集已经被右填充识别的单元格
                right_filled_cells = set()
                for ph in placeholders:
                    if ph.path.startswith(table_path):
                        match = re.search(r'/row[[](\d+)]/cell[[](\d+)]$', ph.path)
                        if match:
                            right_filled_cells.add((int(match.group(1)), int(match.group(2))))

                # 下填充
                for c in range(num_cols):
                    for r in range(1, num_rows):
                        if r in skip_rows or (r - 1) in skip_rows:
                            continue
                        if (r, c) in right_filled_cells: continue
                        current_cell = rows[r].cells[c]
                        top_cell = rows[r - 1].cells[c]

                        if not matrix[r][c].strip() and not is_merged_cell(current_cell) and matrix[r - 1][c].strip():
                            if not is_independent_cell(rows[r], c, current_cell): continue
                            if not is_independent_cell(rows[r - 1], c, top_cell): continue
                            if (r - 1, c) in right_filled_cells: continue

                            is_label_cell = True
                            if c > 0:
                                left_col = c - 1
                                while left_col >= 0 and rows[r - 1].cells[left_col]._element is top_cell._element:
                                    left_col -= 1
                                if left_col >= 0 and matrix[r - 1][left_col].strip():
                                    is_label_cell = False

                            if is_label_cell:
                                path = f"{table_path}/row[{r}]/cell[{c}]"
                                placeholders.append(Placeholder(path=path, label=matrix[r - 1][c], field_key=""))

            return True

        elif tag == 'sdt':
            sdt_content = child.find(qn('w:sdtContent'))
            if sdt_content is not None:
                nonlocal element_index
                for sdt_child in sdt_content:
                    if process_element(sdt_child, element_index):
                        element_index += 1
                return False
        return False

    for child in body:
        if process_element(child, element_index):
            element_index += 1

    for idx, placeholder in enumerate(placeholders, 1):
        placeholder.field_key = f"${idx}"

    return placeholders
