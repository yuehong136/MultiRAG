"""
调试辅助工具：输出文档结构和识别结果
"""

from typing import Any

from docx import Document
from docx.oxml.ns import qn


def extract_document_structure(doc_path: str) -> dict[str, Any]:
    """
    提取文档的完整结构，以可读的方式展示
    """
    doc = Document(doc_path)
    structure = {"paragraphs": [], "tables": []}

    # 遍历所有块级元素
    body = doc.element.body
    element_index = 0

    def process_element(child, index):
        tag = child.tag.split("}")[-1]

        if tag == "p":  # 段落
            para = None
            for p in doc.paragraphs:
                if p._element == child:
                    para = p
                    break

            if para:
                para_info = {"index": index, "path": f"body[{index}]", "type": "paragraph", "text": para.text, "runs": []}

                for run_idx, run in enumerate(para.runs):
                    run_info = {
                        "index": run_idx,
                        "path": f"body[{index}]/run[{run_idx}]",
                        "text": run.text,
                        "bold": bool(run.font.bold),
                        "italic": bool(run.font.italic),
                        "underline": bool(run.font.underline),
                    }
                    para_info["runs"].append(run_info)

                structure["paragraphs"].append(para_info)
                return True

        elif tag == "tbl":  # 表格
            table = None
            for tbl in doc.tables:
                if tbl._element == child:
                    table = tbl
                    break

            if table:
                table_info = {"index": index, "path": f"body[{index}]", "type": "table", "rows": []}

                for row_idx, row in enumerate(table.rows):
                    row_info = {"index": row_idx, "path": f"body[{index}]/row[{row_idx}]", "cells": []}

                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        cell_info = {
                            "index": cell_idx,
                            "path": f"body[{index}]/row[{row_idx}]/cell[{cell_idx}]",
                            "text": cell_text,
                            "text_preview": cell_text[:50] + "..." if len(cell_text) > 50 else cell_text,
                            "paragraphs": [],
                        }

                        for para_idx, para in enumerate(cell.paragraphs):
                            para_info = {"index": para_idx, "path": f"body[{index}]/row[{row_idx}]/cell[{cell_idx}]/p[{para_idx}]", "text": para.text, "runs": []}

                            for run_idx, run in enumerate(para.runs):
                                run_info = {
                                    "index": run_idx,
                                    "path": f"body[{index}]/row[{row_idx}]/cell[{cell_idx}]/p[{para_idx}]/run[{run_idx}]",
                                    "text": run.text,
                                    "underline": bool(run.font.underline),
                                }
                                para_info["runs"].append(run_info)

                            cell_info["paragraphs"].append(para_info)

                        row_info["cells"].append(cell_info)

                    table_info["rows"].append(row_info)

                structure["tables"].append(table_info)
                return True

        elif tag == "sdt":
            sdt_content = child.find(qn("w:sdtContent"))
            if sdt_content is not None:
                nonlocal element_index
                for sdt_child in sdt_content:
                    if process_element(sdt_child, element_index):
                        element_index += 1
                return False

        return False

    for child in body:
        if process_element(child, element_index):
            element_index += 1

    return structure


def format_structure_for_display(structure: dict[str, Any]) -> str:
    """
    将文档结构格式化为易读的文本
    """
    lines = []
    lines.append("=" * 80)
    lines.append("文档结构")
    lines.append("=" * 80)
    lines.append("")

    # 段落
    if structure["paragraphs"]:
        lines.append("【段落】")
        lines.append("-" * 80)
        sorted_paras = sorted(structure["paragraphs"], key=lambda x: x["index"])
        for para in sorted_paras:
            lines.append(f"\n{para['path']}")
            lines.append(f"  文本: {para['text']}")
            if para["runs"]:
                lines.append(f"  Runs ({len(para['runs'])}个):")
                for run in para["runs"]:
                    attrs = []
                    if run["bold"]:
                        attrs.append("粗体")
                    if run["italic"]:
                        attrs.append("斜体")
                    if run["underline"]:
                        attrs.append("下划线")
                    attr_str = f" [{', '.join(attrs)}]" if attrs else ""
                    lines.append(f'    - {run["path"]}: "{run["text"]}"{attr_str}')
        lines.append("")

    # 表格
    if structure["tables"]:
        lines.append("【表格】")
        lines.append("-" * 80)
        sorted_tables = sorted(structure["tables"], key=lambda x: x["index"])
        for table in sorted_tables:
            lines.append(f"\n{table['path']} ({len(table['rows'])}行)")
            for row in table["rows"]:
                lines.append(f"\n  {row['path']} ({len(row['cells'])}列)")
                for cell in row["cells"]:
                    lines.append(f"    {cell['path']}")
                    lines.append(f"      内容: {cell['text_preview']}")
                    if len(cell["paragraphs"]) > 1 or (cell["paragraphs"] and cell["paragraphs"][0]["runs"]):
                        lines.append(f"      段落数: {len(cell['paragraphs'])}")
                        for para in cell["paragraphs"]:
                            if para["runs"]:
                                lines.append(f'        {para["path"]}: "{para["text"]}"')
                                for run in para["runs"]:
                                    if run["underline"]:
                                        lines.append(f'          - {run["path"]}: "{run["text"]}" [下划线]')
        lines.append("")

    lines.append("=" * 80)
    return "\n".join(lines)


def format_placeholders_for_display(placeholders: list[dict[str, Any]]) -> str:
    lines = []
    lines.append("=" * 80)
    lines.append(f"自动识别结果 (共 {len(placeholders)} 个)")
    lines.append("=" * 80)
    lines.append("")

    if not placeholders:
        lines.append("未识别到任何待填项")
    else:
        for idx, ph in enumerate(placeholders, 1):
            lines.append(f"{idx}. {ph['path']}")
            lines.append(f"   标签: {ph['label']}")
            lines.append(f"   字段: {ph.get('field_key', 'N/A')}")
            lines.append("")

    lines.append("=" * 80)
    return "\n".join(lines)


def generate_debug_report(doc_path: str, placeholders: list[dict[str, Any]]) -> str:
    structure = extract_document_structure(doc_path)
    report = []
    report.append(format_structure_for_display(structure))
    report.append("\n\n")
    report.append(format_placeholders_for_display(placeholders))
    return "\n".join(report)
