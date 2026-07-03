"""
DOCX 文档解析器，将文档转换为带路径的结构化数据
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table as DocxTable

from .models import Cell, DocumentElement, Paragraph, ParsedDocument, Row, Run, RunStyle, Table


def get_run_color(run) -> str | None:
    """获取 run 的字体颜色"""
    if run.font.color and run.font.color.rgb:
        return str(run.font.color.rgb)
    return None


def get_alignment_str(alignment) -> str | None:
    """将对齐枚举转换为字符串"""
    if alignment is None:
        return None
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(alignment)


def parse_run(run, path: str) -> Run:
    """解析单个 run"""
    font = run.font
    style = RunStyle(
        bold=bool(font.bold),
        italic=bool(font.italic),
        underline=bool(font.underline),
        font_size=font.size.pt if font.size else None,
        font_name=font.name,
        color=get_run_color(run),
    )
    return Run(path=path, text=run.text, style=style)


def parse_paragraph(para, path: str) -> Paragraph:
    """解析段落"""
    runs = []
    for i, run in enumerate(para.runs):
        run_path = f"{path}/run[{i}]"
        runs.append(parse_run(run, run_path))

    return Paragraph(
        path=path,
        runs=runs,
        alignment=get_alignment_str(para.alignment),
    )


def parse_cell(cell, path: str, row_span: int = 1, col_span: int = 1, is_merged_origin: bool = True, merge_origin_path: str | None = None) -> Cell:
    """解析表格单元格"""
    paragraphs = []
    for i, para in enumerate(cell.paragraphs):
        para_path = f"{path}/p[{i}]"
        paragraphs.append(parse_paragraph(para, para_path))

    return Cell(
        path=path,
        paragraphs=paragraphs,
        row_span=row_span,
        col_span=col_span,
        is_merged_origin=is_merged_origin,
        merge_origin_path=merge_origin_path,
    )


def get_cell_merge_info(table: DocxTable) -> dict:
    """
    分析表格的合并单元格信息

    返回: {(row_idx, col_idx): {
        'row_span': int,
        'col_span': int,
        'is_origin': bool,
        'origin': (origin_row, origin_col) or None
    }}
    """
    merge_info = {}
    num_rows = len(table.rows)

    if num_rows == 0:
        return merge_info

    # 先收集所有行的 cells（只访问一次 table.rows，避免重复创建对象）
    all_rows_cells = [list(row.cells) for row in table.rows]

    # 用 _tc 元素本身作为 key（不用 id()，因为同一个 XML 元素对象是相等的）
    tc_first_pos = {}  # tc element -> (first_row, first_col)

    # 第一遍：记录每个 tc 第一次出现的位置
    for row_idx, row_cells in enumerate(all_rows_cells):
        for col_idx, cell in enumerate(row_cells):
            tc = cell._tc
            if tc not in tc_first_pos:
                tc_first_pos[tc] = (row_idx, col_idx)

    # 第二遍：分析每个单元格
    for row_idx, row_cells in enumerate(all_rows_cells):
        for col_idx, cell in enumerate(row_cells):
            tc = cell._tc
            origin = tc_first_pos[tc]

            if origin == (row_idx, col_idx):
                # 这是合并区域的起始单元格，需要计算 span
                row_span = 1
                col_span = 1

                # 计算 col_span（检查同一行后面的单元格）
                for next_col in range(col_idx + 1, len(row_cells)):
                    if all_rows_cells[row_idx][next_col]._tc is tc:
                        col_span += 1
                    else:
                        break

                # 计算 row_span（检查下面行同一列的单元格）
                for next_row in range(row_idx + 1, num_rows):
                    next_row_cells = all_rows_cells[next_row]
                    if col_idx < len(next_row_cells) and next_row_cells[col_idx]._tc is tc:
                        row_span += 1
                    else:
                        break

                merge_info[(row_idx, col_idx)] = {
                    "row_span": row_span,
                    "col_span": col_span,
                    "is_origin": True,
                    "origin": None,
                }
            else:
                # 这是被合并的单元格
                merge_info[(row_idx, col_idx)] = {
                    "row_span": 1,
                    "col_span": 1,
                    "is_origin": False,
                    "origin": origin,
                }

    return merge_info


def parse_table(table: DocxTable, path: str) -> Table:
    """解析表格，包含合并单元格处理"""
    rows = []
    merge_info = get_cell_merge_info(table)

    for i, row in enumerate(table.rows):
        row_path = f"{path}/row[{i}]"
        cells = []

        for j, cell in enumerate(row.cells):
            cell_path = f"{row_path}/cell[{j}]"
            info = merge_info.get(
                (i, j),
                {
                    "row_span": 1,
                    "col_span": 1,
                    "is_origin": True,
                    "origin": None,
                },
            )

            origin_path = None
            if info["origin"]:
                origin_row, origin_col = info["origin"]
                origin_path = f"{path}/row[{origin_row}]/cell[{origin_col}]"

            cells.append(
                parse_cell(
                    cell,
                    cell_path,
                    row_span=info["row_span"],
                    col_span=info["col_span"],
                    is_merged_origin=info["is_origin"],
                    merge_origin_path=origin_path,
                )
            )

        rows.append(Row(path=row_path, cells=cells))

    return Table(path=path, rows=rows)


def parse_docx(file_path: str, filename: str) -> ParsedDocument:
    """
    解析 docx 文件，返回结构化数据

    关键点：python-docx 的 document.element.body 包含所有块级元素，
    按顺序遍历可以保持段落和表格的相对位置
    """
    doc = Document(file_path)
    elements = []

    # 遍历 body 中的所有块级元素
    body = doc.element.body
    element_index = 0

    def process_element(child, index):
        tag = child.tag.split("}")[-1]

        if tag == "p":  # 段落
            path = f"body[{index}]"
            para = None
            for p in doc.paragraphs:
                if p._element == child:
                    para = p
                    break

            if para:
                elements.append(
                    DocumentElement(
                        type="paragraph",
                        path=path,
                        paragraph=parse_paragraph(para, path),
                    )
                )
                return True

        elif tag == "tbl":  # 表格
            path = f"body[{index}]"
            for tbl in doc.tables:
                if tbl._element == child:
                    elements.append(
                        DocumentElement(
                            type="table",
                            path=path,
                            table=parse_table(tbl, path),
                        )
                    )
                    return True

        elif tag == "sdt":  # 内容控制
            # 递归处理 sdtContent
            sdt_content = child.find(qn("w:sdtContent"))
            if sdt_content is not None:
                nonlocal element_index
                for sdt_child in sdt_content:
                    if process_element(sdt_child, element_index):
                        element_index += 1
                return False

        return False

    for child in body:
        if process_element(child, element_index):
            element_index += 1

    return ParsedDocument(filename=filename, elements=elements)
