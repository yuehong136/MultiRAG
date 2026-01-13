# coding=utf-8
"""
根据 path 定位 docx 元素并填充数据
"""

import re
import copy
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_insert_path(path: str) -> tuple[str, str | None]:
    """
    解析插入路径，分离基础路径和插入位置
    """
    if '::' in path:
        base_path, position = path.rsplit('::', 1)
        return base_path, position
    return path, None


def parse_path(path: str) -> list[tuple[str, int]]:
    """
    解析 path 字符串为 (类型, 索引) 列表
    """
    pattern = r'(\w+)\[(\d+)\]'
    matches = re.findall(pattern, path)
    return [(name, int(idx)) for name, idx in matches]


def get_element_by_path(doc: "DocumentType", path: str):
    """
    根据 path 获取 docx 中对应的元素
    """
    parts = parse_path(path)
    if not parts:
        return None

    # 获取 body 中的元素（段落或表格）
    body = doc.element.body
    
    # 递归查找逻辑，与 parser.py 保持一致
    def find_in_body(target_idx):
        current_idx = 0
        
        def search(container):
            nonlocal current_idx
            for child in container:
                tag = child.tag.split('}')[-1]
                if tag in ('p', 'tbl'):
                    if current_idx == target_idx:
                        return child
                    current_idx += 1
                elif tag == 'sdt':
                    sdt_content = child.find(qn('w:sdtContent'))
                    if sdt_content is not None:
                        res = search(sdt_content)
                        if res is not None:
                            return res
            return None
        
        return search(body)

    # body[n] - 找到第 n 个块级元素
    body_idx = parts[0][1]
    child_element = find_in_body(body_idx)
    
    if child_element is None:
        return None

    # 转换 XML 元素为 python-docx 对象
    current_element = None
    tag = child_element.tag.split('}')[-1]
    if tag == 'p':
        for p in doc.paragraphs:
            if p._element == child_element:
                current_element = p
                break
    elif tag == 'tbl':
        for t in doc.tables:
            if t._element == child_element:
                current_element = t
                break

    if current_element is None:
        return None

    # 只有 body[n]，返回段落或表格
    if len(parts) == 1:
        return current_element

    # 处理后续路径
    for i, (part_type, part_idx) in enumerate(parts[1:], 1):
        if part_type == 'run':
            if hasattr(current_element, 'runs'):
                if part_idx < len(current_element.runs):
                    current_element = current_element.runs[part_idx]
                else:
                    return None
            else:
                return None

        elif part_type == 'row':
            if hasattr(current_element, 'rows'):
                rows = list(current_element.rows)
                if part_idx < len(rows):
                    current_element = rows[part_idx]
                else:
                    return None
            else:
                return None

        elif part_type == 'cell':
            if hasattr(current_element, 'cells'):
                cells = list(current_element.cells)
                if part_idx < len(cells):
                    current_element = cells[part_idx]
                else:
                    return None
            else:
                return None

        elif part_type == 'p':
            if hasattr(current_element, 'paragraphs'):
                if part_idx < len(current_element.paragraphs):
                    current_element = current_element.paragraphs[part_idx]
                else:
                    return None
            else:
                return None

    return current_element


def fill_element(element, value: str):
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

    if isinstance(element, Run):
        element.text = value
    elif isinstance(element, Paragraph):
        if element.runs:
            first_run = element.runs[0]
            for run in element.runs[1:]:
                run._element.getparent().remove(run._element)
            first_run.text = value
        else:
            element.add_run(value)
    elif isinstance(element, _Cell):
        if element.paragraphs:
            para = element.paragraphs[0]
            if para.runs:
                para.runs[0].text = value
                for run in para.runs[1:]:
                    run._element.getparent().remove(run._element)
            else:
                para.add_run(value)
            for p in element.paragraphs[1:]:
                p._element.getparent().remove(p._element)


def insert_run_at(paragraph, ref_run_index: int, position: str, text: str):
    new_run = paragraph.add_run(text)
    run_element = new_run._element
    runs = paragraph.runs
    if position == 'before':
        if ref_run_index < len(runs) - 1:
            ref_element = runs[ref_run_index]._element
            ref_element.addprevious(run_element)
    elif position == 'after':
        if ref_run_index < len(runs) - 1:
            ref_element = runs[ref_run_index]._element
            ref_element.addnext(run_element)


def insert_paragraph_at(cell, ref_para_index: int, position: str, text: str):
    new_para = cell.add_paragraph(text)
    para_element = new_para._element
    paragraphs = cell.paragraphs
    num_paras = len(paragraphs)
    if position == 'prepend':
        if num_paras > 1:
            first_element = paragraphs[0]._element
            first_element.addprevious(para_element)
    elif position == 'before':
        if ref_para_index < num_paras - 1:
            ref_element = paragraphs[ref_para_index]._element
            ref_element.addprevious(para_element)
    elif position == 'after':
        if ref_para_index < num_paras - 1:
            ref_element = paragraphs[ref_para_index]._element
            ref_element.addnext(para_element)


def get_parent_element_by_path(doc: "DocumentType", path: str):
    parts = parse_path(path)
    if len(parts) < 2:
        return None, None, None
    last_part = parts[-1]
    parent_path = '/'.join(f"{name}[{idx}]" for name, idx in parts[:-1])
    parent_element = get_element_by_path(doc, parent_path)
    if parent_element is None:
        return None, None, None
    return parent_element, last_part[0], last_part[1]


def fill_document(doc_path: str, placeholders: list[dict], data: dict, output_path: str) -> str:
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph
    doc = Document(doc_path)
    sorted_placeholders = sorted(placeholders, key=lambda p: p['path'], reverse=True)

    for placeholder in sorted_placeholders:
        placeholder_id = placeholder.get('field_key') or placeholder.get('id')
        if not placeholder_id or placeholder_id not in data:
            continue
        value = data[placeholder_id]
        full_path = placeholder['path']
        base_path, insert_position = parse_insert_path(full_path)

        if insert_position is None:
            element = get_element_by_path(doc, base_path)
            if element:
                fill_element(element, str(value))
        else:
            if insert_position in ('prepend', 'append'):
                container = get_element_by_path(doc, base_path)
                if container and isinstance(container, _Cell):
                    insert_paragraph_at(container, 0, insert_position, str(value))
            else:
                parent, elem_type, elem_index = get_parent_element_by_path(doc, base_path)
                if parent is None: continue
                if elem_type == 'run' and isinstance(parent, Paragraph):
                    insert_run_at(parent, elem_index, insert_position, str(value))
                elif elem_type == 'p' and isinstance(parent, _Cell):
                    insert_paragraph_at(parent, elem_index, insert_position, str(value))

    doc.save(output_path)
    return output_path


def get_table_by_path(doc: "DocumentType", path: str):
    """
    根据 path 获取表格对象
    path 格式: body[n]/row[m]/cell[k] 或 body[n]
    返回表格对象和表格在 body 中的索引
    """
    parts = parse_path(path)
    if not parts:
        return None, None

    # 获取 body 索引
    body_idx = parts[0][1]

    # 获取 body 中的元素
    body = doc.element.body

    def find_in_body(target_idx):
        current_idx = 0

        def search(container):
            nonlocal current_idx
            for child in container:
                tag = child.tag.split('}')[-1]
                if tag in ('p', 'tbl'):
                    if current_idx == target_idx:
                        return child
                    current_idx += 1
                elif tag == 'sdt':
                    sdt_content = child.find(qn('w:sdtContent'))
                    if sdt_content is not None:
                        res = search(sdt_content)
                        if res is not None:
                            return res
            return None

        return search(body)

    child_element = find_in_body(body_idx)
    if child_element is None:
        return None, None

    tag = child_element.tag.split('}')[-1]
    if tag != 'tbl':
        return None, None

    # 找到对应的 python-docx Table 对象
    for t in doc.tables:
        if t._element == child_element:
            return t, body_idx

    return None, None


def copy_row(table, source_row_idx: int) -> None:
    """
    复制表格中的一行并插入到源行之后
    """
    rows = list(table.rows)
    if source_row_idx >= len(rows):
        return

    source_row = rows[source_row_idx]
    # 深拷贝行的 XML 元素
    new_tr = copy.deepcopy(source_row._tr)
    # 在源行后插入新行
    source_row._tr.addnext(new_tr)


def fill_table_cell(table, row_idx: int, cell_idx: int, value: str):
    """
    填充表格中指定位置的单元格
    """
    rows = list(table.rows)
    if row_idx >= len(rows):
        return False

    row = rows[row_idx]
    cells = list(row.cells)
    if cell_idx >= len(cells):
        return False

    cell = cells[cell_idx]
    # 填充单元格
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = value
            for run in para.runs[1:]:
                run._element.getparent().remove(run._element)
        else:
            para.add_run(value)
        # 删除多余的段落
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
    return True


def fill_table_placeholder(doc: "DocumentType", placeholder: dict, rows_data: list[dict[str, str]], row_offset: int = 0) -> int:
    """
    填充表格类型的待填项

    参数:
    - doc: Document 对象
    - placeholder: 待填项配置，包含 path, table_config 等
    - rows_data: 行数据列表，每行是 {列名: 值} 的字典
    - row_offset: 行偏移量（由于之前的动态表格扩展导致的偏移）

    返回:
    - 新增的行数（用于计算后续 placeholder 的偏移量）
    """
    table_config = placeholder.get('table_config')
    if not table_config:
        return 0

    # 获取表格
    table, _ = get_table_by_path(doc, placeholder['path'])
    if table is None:
        return 0

    columns = table_config.get('columns', [])
    data_start_row = table_config.get('data_start_row', 0) + row_offset
    data_end_row = table_config.get('data_end_row', data_start_row) + row_offset
    is_dynamic = table_config.get('dynamic', False)

    if not columns or not rows_data:
        return 0

    # 构建列名到 cell_index 的映射
    column_map = {col['name']: col['cell_index'] for col in columns}

    # 计算现有数据行数
    existing_data_rows = data_end_row - data_start_row + 1

    # 新增的行数
    rows_added = 0

    # 动态表格：如果数据行数超过现有行数，需要复制行
    if is_dynamic and len(rows_data) > existing_data_rows:
        # 需要添加的行数
        rows_to_add = len(rows_data) - existing_data_rows
        rows_added = rows_to_add
        # 从模板行（数据起始行）复制
        for _ in range(rows_to_add):
            copy_row(table, data_start_row)

    # 填充数据
    for row_idx, row_data in enumerate(rows_data):
        actual_row_idx = data_start_row + row_idx

        # 普通表格：不超过 data_end_row
        if not is_dynamic and actual_row_idx > data_end_row:
            break

        # 填充每一列
        for col_name, value in row_data.items():
            if col_name in column_map:
                cell_idx = column_map[col_name]
                fill_table_cell(table, actual_row_idx, cell_idx, value)

    return rows_added


def get_row_index_from_path(path: str) -> int | None:
    """
    从 path 中提取 row 索引
    例如: body[10]/row[4]/cell[0] -> 4
    """
    parts = parse_path(path)
    for name, idx in parts:
        if name == 'row':
            return idx
    return None


def get_body_index_from_path(path: str) -> int | None:
    """
    从 path 中提取 body 索引
    例如: body[10]/row[4]/cell[0] -> 10
    """
    parts = parse_path(path)
    if parts and parts[0][0] == 'body':
        return parts[0][1]
    return None


def apply_row_offset_to_path(path: str, offset: int) -> str:
    """
    将行偏移量应用到 path 中
    例如: body[10]/row[4]/cell[0] + offset=2 -> body[10]/row[6]/cell[0]
    """
    if offset == 0:
        return path

    parts = parse_path(path)
    new_parts = []
    for name, idx in parts:
        if name == 'row':
            new_parts.append(f"{name}[{idx + offset}]")
        else:
            new_parts.append(f"{name}[{idx}]")
    return '/'.join(new_parts)


def fill_document_with_tables(doc_path: str, placeholders: list[dict], data: dict, output_path: str) -> str:
    """
    填充文档（支持表格类型）

    参数:
    - doc_path: 源文档路径
    - placeholders: 待填项列表
    - data: 填充数据，格式为 {field_key: value} 或 {field_key: {"rows": [...]}}
    - output_path: 输出文档路径

    处理逻辑:
    1. 按 body 索引分组所有 placeholder
    2. 在每个 body（表格）内，按 row 索引从小到大排序
    3. 先处理动态表格类型，记录新增的行数
    4. 处理普通类型时，根据累计的行偏移量调整 path
    """
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph

    doc = Document(doc_path)

    # 按 body 索引分组
    body_groups: dict[int, list[dict]] = {}
    non_table_placeholders = []  # 不在表格中的 placeholder（没有 row）

    for placeholder in placeholders:
        body_idx = get_body_index_from_path(placeholder['path'])
        row_idx = get_row_index_from_path(placeholder['path'])

        if body_idx is not None and row_idx is not None:
            if body_idx not in body_groups:
                body_groups[body_idx] = []
            body_groups[body_idx].append(placeholder)
        else:
            non_table_placeholders.append(placeholder)

    # 处理每个 body 组
    for body_idx in sorted(body_groups.keys()):
        group = body_groups[body_idx]

        # 分离表格类型和普通类型
        table_phs = []
        normal_phs = []
        for ph in group:
            ph_type = ph.get('type', 'cell')
            if ph_type in ('table', 'dynamic_table'):
                table_phs.append(ph)
            else:
                normal_phs.append(ph)

        # 按 data_start_row 排序表格类型（从小到大）
        table_phs.sort(key=lambda p: p.get('table_config', {}).get('data_start_row', 0))

        # 按 row 索引排序普通类型（从小到大，用于后续应用偏移量）
        normal_phs.sort(key=lambda p: get_row_index_from_path(p['path']) or 0)

        # 记录每个位置的累计偏移量
        # key: 原始 row 索引, value: 在该位置之前累计新增的行数
        row_offsets: list[tuple[int, int]] = []  # [(data_end_row, rows_added), ...]

        # 先处理表格类型，记录偏移量
        cumulative_offset = 0
        for placeholder in table_phs:
            placeholder_id = placeholder.get('field_key') or placeholder.get('id')
            if not placeholder_id or placeholder_id not in data:
                continue

            field_data = data[placeholder_id]
            if isinstance(field_data, dict) and 'rows' in field_data:
                rows_data = field_data['rows']
            elif isinstance(field_data, list):
                rows_data = field_data
            else:
                continue

            table_config = placeholder.get('table_config', {})
            original_data_end_row = table_config.get('data_end_row', 0)

            # 填充表格，传入当前累计偏移量
            rows_added = fill_table_placeholder(doc, placeholder, rows_data, cumulative_offset)

            # 记录这个表格区域结束后的偏移量
            if rows_added > 0:
                row_offsets.append((original_data_end_row, cumulative_offset + rows_added))
                cumulative_offset += rows_added

        # 处理普通类型，应用偏移量
        for placeholder in normal_phs:
            placeholder_id = placeholder.get('field_key') or placeholder.get('id')
            if not placeholder_id or placeholder_id not in data:
                continue

            field_data = data[placeholder_id]
            if isinstance(field_data, dict):
                value = field_data.get('value', '')
            else:
                value = str(field_data)

            # 计算该 placeholder 应该应用的偏移量
            original_row_idx = get_row_index_from_path(placeholder['path'])
            offset_to_apply = 0
            if original_row_idx is not None:
                for data_end_row, offset in row_offsets:
                    if original_row_idx > data_end_row:
                        offset_to_apply = offset

            # 应用偏移量到 path
            full_path = placeholder['path']
            adjusted_path = apply_row_offset_to_path(full_path, offset_to_apply)
            base_path, insert_position = parse_insert_path(adjusted_path)

            if insert_position is None:
                element = get_element_by_path(doc, base_path)
                if element:
                    fill_element(element, str(value))
            else:
                if insert_position in ('prepend', 'append'):
                    container = get_element_by_path(doc, base_path)
                    if container and isinstance(container, _Cell):
                        insert_paragraph_at(container, 0, insert_position, str(value))
                else:
                    parent, elem_type, elem_index = get_parent_element_by_path(doc, base_path)
                    if parent is None:
                        continue
                    if elem_type == 'run' and isinstance(parent, Paragraph):
                        insert_run_at(parent, elem_index, insert_position, str(value))
                    elif elem_type == 'p' and isinstance(parent, _Cell):
                        insert_paragraph_at(parent, elem_index, insert_position, str(value))

    # 处理不在表格中的 placeholder（没有 row 的）
    for placeholder in non_table_placeholders:
        placeholder_id = placeholder.get('field_key') or placeholder.get('id')
        if not placeholder_id or placeholder_id not in data:
            continue

        field_data = data[placeholder_id]
        if isinstance(field_data, dict):
            value = field_data.get('value', '')
        else:
            value = str(field_data)

        full_path = placeholder['path']
        base_path, insert_position = parse_insert_path(full_path)

        if insert_position is None:
            element = get_element_by_path(doc, base_path)
            if element:
                fill_element(element, str(value))
        else:
            if insert_position in ('prepend', 'append'):
                container = get_element_by_path(doc, base_path)
                if container and isinstance(container, _Cell):
                    insert_paragraph_at(container, 0, insert_position, str(value))
            else:
                parent, elem_type, elem_index = get_parent_element_by_path(doc, base_path)
                if parent is None:
                    continue
                if elem_type == 'run' and isinstance(parent, Paragraph):
                    insert_run_at(parent, elem_index, insert_position, str(value))
                elif elem_type == 'p' and isinstance(parent, _Cell):
                    insert_paragraph_at(parent, elem_index, insert_position, str(value))

    doc.save(output_path)
    return output_path